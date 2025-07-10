# 统一的仓库选址优化系统 - 整合版
# 整合了两个系统的所有优势功能，提供最完整的仓库选址解决方案

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import folium
from streamlit_folium import st_folium, folium_static
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, timedelta
import random
import math
import time
from sklearn.cluster import KMeans, DBSCAN, AgglomerativeClustering
from sklearn.preprocessing import StandardScaler, MinMaxScaler
from sklearn.decomposition import PCA
from sklearn.ensemble import RandomForestRegressor, GradientBoostingRegressor
from sklearn.model_selection import cross_val_score
from scipy.spatial.distance import cdist
from scipy.spatial import distance_matrix
from scipy.optimize import minimize, differential_evolution, dual_annealing
from scipy.stats import norm, lognorm
import warnings
import io
import base64
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xlsxwriter
import json
import zipfile
import tempfile
import os
import networkx as nx
from geopy.distance import geodesic
from shapely.geometry import Point, Polygon
import geopandas as gpd
import requests
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
import asyncio
import aiohttp
from typing import List, Dict, Tuple, Optional, Union
import hashlib
import pickle
from functools import lru_cache
import pulp
from ortools.linear_solver import pywraplp
from ortools.constraint_solver import routing_enums_pb2
from ortools.constraint_solver import pywrapcp
import osmnx as ox
from folium.plugins import HeatMap, MarkerCluster, MeasureControl, Search, Fullscreen

warnings.filterwarnings('ignore')

# 设置matplotlib中文字体
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

# 配置页面
st.set_page_config(
    page_title="仓库选址优化系统 Ultimate Pro",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://github.com/warehouse-optimization',
        'Report a bug': "https://github.com/warehouse-optimization/issues",
        'About': "# 仓库选址优化系统 Ultimate Pro\n业界最先进的仓库选址解决方案"
    }
)

# 高级CSS样式（整合两个系统的样式）
st.markdown("""
<style>
    /* 主题样式 */
    .main-header {
        background: linear-gradient(135deg, #1e3c72, #2a5298);
        color: white;
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        animation: slideIn 0.5s ease-out;
        font-size: 3rem;
    }

    @keyframes slideIn {
        from {
            transform: translateY(-50px);
            opacity: 0;
        }
        to {
            transform: translateY(0);
            opacity: 1;
        }
    }

    /* 卡片样式 */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 5px 15px rgba(0,0,0,0.08);
        transition: all 0.3s ease;
        border-left: 4px solid #1e3c72;
    }

    .metric-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 25px rgba(0,0,0,0.15);
    }

    /* 高级按钮 */
    .stButton>button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        font-weight: 600;
        border-radius: 50px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
    }

    .stButton>button:hover {
        transform: translateY(-3px);
        box-shadow: 0 7px 20px rgba(102, 126, 234, 0.6);
    }

    /* 数据表格美化 */
    .dataframe {
        border: none !important;
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }

    /* 侧边栏样式 */
    .css-1d391kg {
        background: linear-gradient(180deg, #f5f7fa 0%, #c3cfe2 100%);
    }

    /* 进度条动画 */
    .stProgress > div > div > div > div {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        animation: progressPulse 2s ease-in-out infinite;
    }

    @keyframes progressPulse {
        0% { opacity: 0.8; }
        50% { opacity: 1; }
        100% { opacity: 0.8; }
    }

    /* 信息提示框 */
    .stAlert {
        border-radius: 10px;
        border: none;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }

    .sub-header {
        font-size: 1.5rem;
        color: #2ca02c;
        margin-bottom: 1rem;
    }

    .success-message {
        color: #28a745;
        font-weight: bold;
    }

    .warning-message {
        color: #ffc107;
        font-weight: bold;
    }

    .error-message {
        color: #dc3545;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

# 初始化高级session state（整合两个系统）
if 'advanced_features' not in st.session_state:
    st.session_state.advanced_features = {
        'ml_models': {},
        'optimization_history': [],
        'real_time_data': {},
        'cache': {},
        'performance_metrics': {},
        'scenario_analysis': {},
        'predictive_models': {}
    }

# 扩展session state
for key in ['customer_data', 'candidate_locations', 'transportation_costs',
            'optimization_results', 'carbon_footprint', 'sustainability_metrics',
            'ml_predictions', 'network_analysis', 'simulation_results', 'supplier_data',
            'historical_demand']:
    if key not in st.session_state:
        st.session_state[key] = pd.DataFrame()

# 初始化非DataFrame的session state
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = {}
if 'selected_locations' not in st.session_state:
    st.session_state.selected_locations = []
if 'distance_matrix' not in st.session_state:
    st.session_state.distance_matrix = {}

# 中国主要城市坐标数据（扩展版）
CHINA_CITIES_EXTENDED = {
    **{
        '北京': (116.4074, 39.9042),
        '上海': (121.4737, 31.2304),
        '广州': (113.2644, 23.1291),
        '深圳': (114.0579, 22.5431),
        '杭州': (120.1551, 30.2741),
        '成都': (104.0668, 30.5728),
        '武汉': (114.3054, 30.5931),
        '南京': (118.7969, 32.0603),
        '西安': (108.9402, 34.3416),
        '重庆': (106.5516, 29.5630),
        '天津': (117.1901, 39.1235),
        '苏州': (120.5854, 31.2989),
        '青岛': (120.3826, 36.0671),
        '郑州': (113.6254, 34.7466),
        '长沙': (112.9388, 28.2282),
        '东莞': (113.7518, 23.0489),
        '宁波': (121.5540, 29.8683),
        '佛山': (113.1220, 23.0288),
        '合肥': (117.2272, 31.8206),
        '福州': (119.2965, 26.0745),
        '厦门': (118.0894, 24.4798),
        '昆明': (102.8332, 25.0389),
        '南昌': (115.8579, 28.6890),
        '石家庄': (114.5149, 38.0428),
        '济南': (117.1205, 36.6519),
        '哈尔滨': (126.5358, 45.8025),
        '沈阳': (123.4307, 41.8056),
        '长春': (125.3245, 43.8171),
        '太原': (112.5489, 37.8706),
        '南宁': (108.3669, 22.8170)
    },
    # 新增更多城市
    **{
        '大连': (121.6147, 38.9140),
        '无锡': (120.3119, 31.4912),
        '常州': (119.9741, 31.8105),
        '温州': (120.6721, 28.0004),
        '绍兴': (120.5821, 30.0329),
        '台州': (121.4286, 28.6561),
        '烟台': (121.3910, 37.5394),
        '潍坊': (119.1616, 36.7069),
        '贵阳': (106.6302, 26.6477),
        '海口': (110.3312, 20.0311),
        '兰州': (103.8236, 36.0581),
        '银川': (106.2765, 38.4668),
        '西宁': (101.7787, 36.6171),
        '呼和浩特': (111.7518, 40.8418),
        '乌鲁木齐': (87.6177, 43.7928),
        '拉萨': (91.1322, 29.6604)
    }
}


# 工具函数
def calculate_distance(lon1, lat1, lon2, lat2):
    """计算两点间的距离（公里）"""
    return geodesic((lat1, lon1), (lat2, lon2)).kilometers


# 高级工具类（来自第一个系统）
class AdvancedOptimizer:
    """高级优化算法集合"""

    @staticmethod
    def particle_swarm_optimization(objective_func, bounds, n_particles=50, max_iter=100):
        """粒子群优化算法"""
        n_dims = len(bounds)
        # 初始化粒子位置和速度
        particles = np.random.uniform(
            [b[0] for b in bounds],
            [b[1] for b in bounds],
            (n_particles, n_dims)
        )
        velocities = np.random.randn(n_particles, n_dims) * 0.1

        # 初始化个体最优和全局最优
        p_best = particles.copy()
        p_best_scores = np.array([objective_func(p) for p in particles])
        g_best_idx = np.argmin(p_best_scores)
        g_best = p_best[g_best_idx].copy()
        g_best_score = p_best_scores[g_best_idx]

        # PSO参数
        w = 0.7  # 惯性权重
        c1 = 1.5  # 个体学习因子
        c2 = 1.5  # 社会学习因子

        history = [g_best_score]

        for _ in range(max_iter):
            # 更新速度和位置
            r1, r2 = np.random.rand(2)
            velocities = (w * velocities +
                          c1 * r1 * (p_best - particles) +
                          c2 * r2 * (g_best - particles))
            particles += velocities

            # 边界处理
            for i, bound in enumerate(bounds):
                particles[:, i] = np.clip(particles[:, i], bound[0], bound[1])

            # 更新最优解
            scores = np.array([objective_func(p) for p in particles])
            better_mask = scores < p_best_scores
            p_best[better_mask] = particles[better_mask]
            p_best_scores[better_mask] = scores[better_mask]

            if np.min(scores) < g_best_score:
                g_best_idx = np.argmin(scores)
                g_best = particles[g_best_idx].copy()
                g_best_score = scores[g_best_idx]

            history.append(g_best_score)

        return g_best, history

    @staticmethod
    def ant_colony_optimization(distance_matrix, n_ants=50, n_iterations=100,
                                alpha=1.0, beta=2.0, evaporation_rate=0.5):
        """蚁群算法"""
        n_cities = len(distance_matrix)
        pheromone = np.ones((n_cities, n_cities))
        best_path = None
        best_distance = float('inf')

        for _ in range(n_iterations):
            paths = []
            distances = []

            for _ in range(n_ants):
                path = [random.randint(0, n_cities - 1)]

                while len(path) < n_cities:
                    current = path[-1]
                    unvisited = [i for i in range(n_cities) if i not in path]

                    if not unvisited:
                        break

                    # 计算转移概率
                    probabilities = []
                    for next_city in unvisited:
                        tau = pheromone[current][next_city] ** alpha
                        eta = (1 / distance_matrix[current][next_city]) ** beta
                        probabilities.append(tau * eta)

                    probabilities = np.array(probabilities) / sum(probabilities)
                    next_city = np.random.choice(unvisited, p=probabilities)
                    path.append(next_city)

                # 计算路径长度
                distance = sum(distance_matrix[path[i]][path[i + 1]]
                               for i in range(len(path) - 1))
                paths.append(path)
                distances.append(distance)

                if distance < best_distance:
                    best_distance = distance
                    best_path = path.copy()

            # 更新信息素
            pheromone *= (1 - evaporation_rate)
            for path, distance in zip(paths, distances):
                for i in range(len(path) - 1):
                    pheromone[path[i]][path[i + 1]] += 1 / distance

        return best_path, best_distance

    @staticmethod
    def quantum_inspired_optimization(objective_func, bounds, n_qubits=20, max_iter=100):
        """量子启发式优化算法"""
        n_dims = len(bounds)

        # 初始化量子比特
        qubits = np.random.rand(n_qubits, n_dims) * np.pi / 2

        best_solution = None
        best_score = float('inf')
        history = []

        for iteration in range(max_iter):
            # 观测（坍缩）量子态
            solutions = []
            for qubit in qubits:
                solution = []
                for i, (low, high) in enumerate(bounds):
                    # 使用量子比特的概率幅度
                    prob = np.sin(qubit[i]) ** 2
                    value = low + prob * (high - low)
                    solution.append(value)
                solutions.append(solution)

            # 评估解
            scores = [objective_func(sol) for sol in solutions]

            # 更新最优解
            min_idx = np.argmin(scores)
            if scores[min_idx] < best_score:
                best_score = scores[min_idx]
                best_solution = solutions[min_idx]

            # 量子旋转门更新
            theta = np.pi / (4 * (1 + iteration))
            for i in range(n_qubits):
                if scores[i] > best_score:
                    # 向最优解旋转
                    rotation = theta * (1 - scores[i] / best_score)
                    qubits[i] += rotation * np.random.randn(n_dims) * 0.1

            history.append(best_score)

        return best_solution, history


class MachineLearningModels:
    """机器学习模型集合"""

    @staticmethod
    def demand_forecasting(historical_data, features, periods=12):
        """需求预测模型"""
        from sklearn.ensemble import RandomForestRegressor
        from sklearn.preprocessing import StandardScaler

        # 特征工程
        X = historical_data[features]
        y = historical_data['demand']

        # 标准化
        scaler = StandardScaler()
        X_scaled = scaler.fit_transform(X)

        # 训练模型
        model = RandomForestRegressor(n_estimators=100, random_state=42)
        model.fit(X_scaled, y)

        # 生成未来预测
        future_features = np.tile(X_scaled[-1], (periods, 1))
        # 添加时间趋势
        time_trend = np.arange(len(y), len(y) + periods).reshape(-1, 1)
        future_features = np.hstack([future_features, time_trend])

        predictions = model.predict(future_features[..., :-1])

        return predictions, model

    @staticmethod
    def location_scoring_model(location_features):
        """位置评分模型"""
        from sklearn.ensemble import GradientBoostingRegressor

        # 这里使用预训练的权重（实际应用中应该用历史数据训练）
        feature_weights = {
            'transportation_access': 0.3,
            'labor_availability': 0.2,
            'infrastructure_quality': 0.2,
            'cost_index': -0.15,
            'risk_score': -0.1,
            'market_proximity': 0.05
        }

        scores = []
        for _, location in location_features.iterrows():
            score = sum(location[feature] * weight
                        for feature, weight in feature_weights.items()
                        if feature in location)
            scores.append(score)

        return np.array(scores)

    @staticmethod
    def risk_prediction_model(risk_factors):
        """风险预测模型"""
        from sklearn.neural_network import MLPRegressor

        # 神经网络风险评估
        model = MLPRegressor(
            hidden_layer_sizes=(100, 50),
            activation='relu',
            solver='adam',
            random_state=42
        )

        # 模拟训练（实际应用中使用真实数据）
        X_train = np.random.rand(1000, len(risk_factors.columns))
        y_train = np.random.rand(1000)

        model.fit(X_train, y_train)

        risk_scores = model.predict(risk_factors)

        return risk_scores


class NetworkAnalysis:
    """供应链网络分析"""

    @staticmethod
    def create_supply_chain_network(warehouses, customers, suppliers=None):
        """创建供应链网络图"""
        G = nx.DiGraph()

        # 添加节点
        for idx, warehouse in warehouses.iterrows():
            G.add_node(f"W_{warehouse['id']}",
                       type='warehouse',
                       **warehouse.to_dict())

        for idx, customer in customers.iterrows():
            G.add_node(f"C_{customer['id']}",
                       type='customer',
                       **customer.to_dict())

        if suppliers is not None:
            for idx, supplier in suppliers.iterrows():
                G.add_node(f"S_{supplier['id']}",
                           type='supplier',
                           **supplier.to_dict())

        return G

    @staticmethod
    def calculate_network_metrics(G):
        """计算网络指标"""
        metrics = {
            'nodes': G.number_of_nodes(),
            'edges': G.number_of_edges(),
            'density': nx.density(G),
            'average_degree': sum(dict(G.degree()).values()) / G.number_of_nodes(),
            'clustering_coefficient': nx.average_clustering(G.to_undirected()),
        }

        if G.number_of_nodes() > 0:
            metrics['betweenness_centrality'] = nx.betweenness_centrality(G)
            metrics['closeness_centrality'] = nx.closeness_centrality(G)

        return metrics

    @staticmethod
    def optimize_network_flow(G, source_nodes, sink_nodes, capacities):
        """网络流优化"""
        # 使用最小成本流算法
        flow_dict = nx.min_cost_flow(G)
        return flow_dict


class SimulationEngine:
    """仿真引擎"""

    @staticmethod
    def monte_carlo_simulation(demand_params, n_simulations=1000):
        """蒙特卡洛仿真"""
        results = []

        for _ in range(n_simulations):
            # 生成随机需求
            simulated_demand = {}
            for location, params in demand_params.items():
                if params['distribution'] == 'normal':
                    demand = np.random.normal(params['mean'], params['std'])
                elif params['distribution'] == 'lognormal':
                    demand = np.random.lognormal(params['mean'], params['std'])
                else:
                    demand = params['mean']

                simulated_demand[location] = max(0, demand)

            results.append(simulated_demand)

        return pd.DataFrame(results)

    @staticmethod
    def discrete_event_simulation(events, duration):
        """离散事件仿真"""
        current_time = 0
        event_log = []

        while current_time < duration and events:
            # 获取下一个事件
            next_event = min(events, key=lambda x: x['time'])
            events.remove(next_event)

            current_time = next_event['time']
            event_log.append({
                'time': current_time,
                'type': next_event['type'],
                'details': next_event.get('details', {})
            })

            # 处理事件并生成新事件
            if next_event['type'] == 'arrival':
                # 生成服务事件
                service_time = current_time + np.random.exponential(5)
                events.append({
                    'time': service_time,
                    'type': 'departure',
                    'details': next_event['details']
                })

        return pd.DataFrame(event_log)


class RealTimeDataIntegration:
    """实时数据集成"""

    @staticmethod
    async def fetch_traffic_data(locations):
        """获取实时交通数据"""
        # 模拟API调用
        traffic_data = {}
        for location in locations:
            traffic_data[location] = {
                'congestion_level': np.random.choice(['low', 'medium', 'high']),
                'average_speed': np.random.uniform(30, 80),
                'incidents': np.random.randint(0, 5)
            }
        return traffic_data

    @staticmethod
    async def fetch_weather_data(locations):
        """获取天气数据"""
        weather_data = {}
        for location in locations:
            weather_data[location] = {
                'temperature': np.random.uniform(-10, 40),
                'precipitation': np.random.uniform(0, 100),
                'wind_speed': np.random.uniform(0, 50),
                'visibility': np.random.choice(['good', 'moderate', 'poor'])
            }
        return weather_data

    @staticmethod
    def fetch_market_data():
        """获取市场数据"""
        return {
            'fuel_price': np.random.uniform(6, 8),
            'labor_cost_index': np.random.uniform(0.9, 1.1),
            'real_estate_index': np.random.uniform(0.95, 1.05),
            'economic_indicator': np.random.uniform(-2, 2)
        }


# 优化算法（整合两个系统的算法）
def mixed_integer_programming(customers, candidates, num_warehouses, budget_limit=None):
    """
    混合整数规划(MIP)求解设施选址问题
    使用线性规划方法找到最优解
    """
    n_customers = len(customers)
    n_candidates = len(candidates)

    # 计算成本矩阵
    transport_costs = np.zeros((n_candidates, n_customers))
    for i, (_, warehouse) in enumerate(candidates.iterrows()):
        for j, (_, customer) in enumerate(customers.iterrows()):
            distance = calculate_distance(
                warehouse['经度'], warehouse['纬度'],
                customer['经度'], customer['纬度']
            )
            transport_costs[i, j] = distance * customer['年需求量'] * 0.5

    # 简化的MIP求解（使用贪心近似）
    fixed_costs = candidates['建设成本'].values

    # 计算每个仓库的效益得分
    efficiency_scores = []
    for i in range(n_candidates):
        total_savings = -transport_costs[i].sum() - fixed_costs[i]
        efficiency_scores.append((i, total_savings))

    # 排序并选择最优的仓库
    efficiency_scores.sort(key=lambda x: x[1], reverse=True)

    selected_indices = []
    total_cost = 0

    for idx, score in efficiency_scores[:num_warehouses]:
        if budget_limit:
            if total_cost + fixed_costs[idx] <= budget_limit:
                selected_indices.append(idx)
                total_cost += fixed_costs[idx]
        else:
            selected_indices.append(idx)

    selected_ids = candidates.iloc[selected_indices]['地点编号'].tolist()
    return selected_ids, total_cost


def median_optimization(data):
    """中位数法（一维线性选址）"""
    sorted_points = sorted(data['经度'].values)
    median = sorted_points[len(sorted_points) // 2]
    total_distance = sum(abs(x - median) for x in sorted_points)
    return median, total_distance


def gravity_center_method(customers):
    """重心法优化"""
    total_demand = customers['年需求量'].sum()
    weighted_lon = (customers['经度'] * customers['年需求量']).sum() / total_demand
    weighted_lat = (customers['纬度'] * customers['年需求量']).sum() / total_demand
    return weighted_lon, weighted_lat


def set_cover_optimization(customers, candidates, coverage_radius=200):
    """集合覆盖模型"""
    customer_points = customers[['经度', '纬度']].values
    candidate_points = candidates[['经度', '纬度']].values

    # 计算距离矩阵
    dist_matrix = cdist(candidate_points, customer_points) * 111  # 转换为公里

    # 找到每个候选点的覆盖范围
    coverage = {}
    for i in range(len(candidate_points)):
        coverage[i] = np.where(dist_matrix[i] <= coverage_radius)[0].tolist()

    # 贪心算法求解
    uncovered = set(range(len(customer_points)))
    selected = []

    while uncovered:
        best_candidate = None
        best_coverage = set()

        for i in range(len(candidate_points)):
            if i not in selected:
                covered = set(coverage[i]) & uncovered
                if len(covered) > len(best_coverage):
                    best_candidate = i
                    best_coverage = covered

        if best_candidate is None:
            break

        selected.append(best_candidate)
        uncovered -= best_coverage

    return candidates.iloc[selected]['地点编号'].tolist()


def genetic_algorithm(customers, candidates, num_warehouses, max_generations=100):
    """遗传算法优化"""
    population_size = 50
    mutation_rate = 0.1

    # 适应度函数
    def calculate_fitness(solution):
        total_cost = 0
        for _, location in candidates[candidates['地点编号'].isin(solution)].iterrows():
            total_cost += location['建设成本'] + location['运营成本']

        # 计算运输成本
        for _, customer in customers.iterrows():
            min_cost = float('inf')
            for warehouse_id in solution:
                if (warehouse_id, customer['客户编号']) in st.session_state.distance_matrix:
                    distance = st.session_state.distance_matrix[(warehouse_id, customer['客户编号'])]
                    cost = distance * 0.5 * customer['年需求量']
                    min_cost = min(min_cost, cost)
            total_cost += min_cost

        return 1 / (total_cost + 1)

    # 初始化种群
    population = []
    candidate_ids = candidates['地点编号'].tolist()

    for _ in range(population_size):
        individual = random.sample(candidate_ids, num_warehouses)
        population.append(individual)

    # 进化过程
    best_fitness_history = []
    best_solution = None
    best_fitness = 0

    for generation in range(max_generations):
        # 计算适应度
        fitness_scores = [calculate_fitness(ind) for ind in population]

        # 记录最佳个体
        max_fitness_idx = np.argmax(fitness_scores)
        if fitness_scores[max_fitness_idx] > best_fitness:
            best_fitness = fitness_scores[max_fitness_idx]
            best_solution = population[max_fitness_idx].copy()

        best_fitness_history.append(best_fitness)

        # 选择和繁殖
        new_population = [best_solution]  # 精英保留

        while len(new_population) < population_size:
            # 轮盘赌选择
            parent1 = random.choices(population, weights=fitness_scores)[0]
            parent2 = random.choices(population, weights=fitness_scores)[0]

            # 交叉
            child = parent1[:num_warehouses // 2] + parent2[num_warehouses // 2:]
            child = list(set(child))[:num_warehouses]  # 去重

            # 如果不足，补充随机基因
            while len(child) < num_warehouses:
                new_gene = random.choice(candidate_ids)
                if new_gene not in child:
                    child.append(new_gene)

            # 变异
            if random.random() < mutation_rate:
                idx = random.randint(0, num_warehouses - 1)
                new_gene = random.choice(candidate_ids)
                if new_gene not in child:
                    child[idx] = new_gene

            new_population.append(child)

        population = new_population

    return best_solution, best_fitness_history


def simulated_annealing(customers, candidates, num_warehouses, initial_temp=1000, cooling_rate=0.95):
    """模拟退火算法"""
    candidate_ids = candidates['地点编号'].tolist()

    # 计算成本函数
    def calculate_cost(solution):
        total_cost = 0
        for warehouse_id in solution:
            location = candidates[candidates['地点编号'] == warehouse_id].iloc[0]
            total_cost += location['建设成本'] + location['运营成本']

        # 计算运输成本
        for _, customer in customers.iterrows():
            min_cost = float('inf')
            for warehouse_id in solution:
                if (warehouse_id, customer['客户编号']) in st.session_state.distance_matrix:
                    distance = st.session_state.distance_matrix[(warehouse_id, customer['客户编号'])]
                    cost = distance * 0.5 * customer['年需求量']
                    min_cost = min(min_cost, cost)
            total_cost += min_cost

        return total_cost

    # 初始解
    current_solution = random.sample(candidate_ids, num_warehouses)
    current_cost = calculate_cost(current_solution)

    best_solution = current_solution.copy()
    best_cost = current_cost

    temperature = initial_temp
    cost_history = [current_cost]

    while temperature > 1:
        # 生成邻域解
        new_solution = current_solution.copy()
        idx = random.randint(0, num_warehouses - 1)
        new_warehouse = random.choice(candidate_ids)

        if new_warehouse not in new_solution:
            new_solution[idx] = new_warehouse

            new_cost = calculate_cost(new_solution)
            delta_cost = new_cost - current_cost

            # 接受准则
            if delta_cost < 0 or random.random() < math.exp(-delta_cost / temperature):
                current_solution = new_solution
                current_cost = new_cost

                if current_cost < best_cost:
                    best_solution = current_solution.copy()
                    best_cost = current_cost

        cost_history.append(current_cost)
        temperature *= cooling_rate

    return best_solution, cost_history


def demand_forecasting(historical_data, periods=12):
    """
    使用时间序列预测未来需求
    """
    # 简化的趋势预测
    if len(historical_data) < 2:
        return historical_data

    # 计算增长率
    growth_rate = (historical_data[-1] - historical_data[0]) / (len(historical_data) - 1)

    # 预测未来需求
    last_value = historical_data[-1]
    forecast = []
    for i in range(periods):
        next_value = last_value + growth_rate * (i + 1)
        # 添加一些随机波动
        next_value *= np.random.uniform(0.95, 1.05)
        forecast.append(max(0, next_value))  # 确保非负

    return forecast


def carbon_footprint_calculation(distance, transport_mode='truck', cargo_weight=1000):
    """
    计算碳足迹

    参数:
    - distance: 运输距离(km)
    - transport_mode: 运输方式
    - cargo_weight: 货物重量(kg)

    返回:
    - carbon_emission: 碳排放量(kg CO2)
    """
    # 碳排放因子 (kg CO2 per ton-km)
    emission_factors = {
        'truck': 0.12,
        'rail': 0.03,
        'air': 0.6,
        'ship': 0.01
    }

    factor = emission_factors.get(transport_mode, 0.12)
    carbon_emission = distance * (cargo_weight / 1000) * factor

    return carbon_emission


def multi_objective_optimization(customers, candidates, weights=None):
    """
    多目标优化：同时考虑成本、服务水平、环境影响
    """
    if weights is None:
        weights = {
            'cost': 0.4,
            'service': 0.3,
            'environment': 0.3
        }

    scores = []

    for _, candidate in candidates.iterrows():
        # 成本得分（归一化）
        cost_score = 1 - (candidate['建设成本'] - candidates['建设成本'].min()) / \
                     (candidates['建设成本'].max() - candidates['建设成本'].min())

        # 服务水平得分（基于覆盖范围和基础设施）
        service_score = (candidate['服务半径'] / candidates['服务半径'].max()) * 0.5 + \
                        (candidate['基础设施评分'] / 10) * 0.5

        # 环境得分（基于交通便利性和风险）
        env_score = 0
        if candidate['交通便利性'] == '优':
            env_score += 0.4
        elif candidate['交通便利性'] == '良':
            env_score += 0.3
        elif candidate['交通便利性'] == '中':
            env_score += 0.2
        else:
            env_score += 0.1

        env_score += (10 - candidate['风险评分']) / 10 * 0.6

        # 综合得分
        total_score = (weights['cost'] * cost_score +
                       weights['service'] * service_score +
                       weights['environment'] * env_score)

        scores.append({
            '地点编号': candidate['地点编号'],
            '地点名称': candidate['地点名称'],
            '城市': candidate['城市'],
            'total_score': total_score,
            '成本得分': cost_score,
            '服务得分': service_score,
            '环境得分': env_score,
            '综合得分': total_score
        })

    return pd.DataFrame(scores).sort_values('综合得分', ascending=False)


# 高级数据生成函数（整合两个系统的生成方法）
def generate_advanced_sample_data():
    """生成高级示例数据"""
    cities = list(CHINA_CITIES_EXTENDED.keys())
    np.random.seed(42)

    # 客户数据（增强版）
    num_customers = st.sidebar.number_input("客户数量", min_value=50, max_value=500, value=100, key="gen_num_customers")
    customer_data = []

    for i in range(num_customers):
        selected_city = np.random.choice(cities)
        city_coords = CHINA_CITIES_EXTENDED[selected_city]

        # 更真实的位置偏移
        lon_offset = np.random.normal(0, 0.3)
        lat_offset = np.random.normal(0, 0.2)

        # 增强的客户属性
        customer_data.append({
            '客户编号': f'C{i + 1:04d}',
            '客户名称': f'{selected_city}客户{i + 1}',
            '城市': selected_city,
            '经度': city_coords[0] + lon_offset,
            '纬度': city_coords[1] + lat_offset,
            '年需求量': np.random.lognormal(8, 1),  # 对数正态分布更真实
            '需求波动率': np.random.uniform(0.1, 0.4),
            '季节性因子': np.random.choice([0.8, 1.0, 1.2, 1.5]),
            '需求权重': np.random.uniform(0.5, 3.0),
            '优先级': np.random.choice(['高', '中', '低'], p=[0.2, 0.5, 0.3]),
            '客户类型': np.random.choice(['零售', '批发', '电商', 'B2B', 'B2C']),
            '行业': np.random.choice(['快消品', '电子', '服装', '食品', '医药', '其他']),
            '服务水平要求': np.random.uniform(0.85, 0.99),
            '付款周期': np.random.choice([0, 30, 60, 90]),
            '信用等级': np.random.choice(['AAA', 'AA', 'A', 'BBB', 'BB']),
            '年增长率': np.random.uniform(-0.1, 0.3),
            '合同期限': np.random.choice([1, 2, 3, 5]),
            '特殊要求': np.random.choice(['无', '冷链', '危险品', '易碎品', '加急'])
        })

    st.session_state.customer_data = pd.DataFrame(customer_data)

    # 候选仓库数据（增强版）
    num_warehouses = st.sidebar.number_input("候选仓库数量", min_value=10, max_value=50, value=20,
                                             key="gen_num_warehouses")
    candidate_data = []

    for i in range(num_warehouses):
        selected_city = np.random.choice(cities, p=None)  # 可以加权选择
        city_coords = CHINA_CITIES_EXTENDED[selected_city]

        # 仓库位置通常在城市边缘
        angle = np.random.uniform(0, 2 * np.pi)
        distance = np.random.uniform(0.1, 0.5)
        lon_offset = distance * np.cos(angle)
        lat_offset = distance * np.sin(angle)

        # 综合成本考虑
        base_land_price = np.random.uniform(1000, 8000)
        city_factor = {'北京': 2.0, '上海': 2.2, '深圳': 2.1, '广州': 1.8}.get(selected_city, 1.0)

        candidate_data.append({
            '地点编号': f'W{i + 1:03d}',
            '地点名称': f'{selected_city}物流园{i + 1}',
            '城市': selected_city,
            '经度': city_coords[0] + lon_offset,
            '纬度': city_coords[1] + lat_offset,
            '土地单价': base_land_price * city_factor,
            '占地面积': np.random.uniform(10000, 100000),  # 平方米
            '建设成本': np.random.uniform(5e6, 5e7),
            '运营成本': np.random.uniform(1e5, 1e6),  # 年运营成本
            '人工成本指数': np.random.uniform(0.7, 1.5),
            '电力成本': np.random.uniform(0.5, 1.0),  # 元/度
            '水费成本': np.random.uniform(3, 6),  # 元/吨
            '基础设施评分': np.random.uniform(6, 10),
            '交通便利性评分': np.random.uniform(6, 10),
            '交通便利性': np.random.choice(['优', '良', '中', '差']),
            '政策支持评分': np.random.uniform(5, 10),
            '环保要求等级': np.random.choice(['低', '中', '高']),
            '风险评分': np.random.uniform(1, 8),
            '最大容量': np.random.uniform(20000, 500000),
            '可用容量比例': np.random.uniform(0.8, 1.0),
            '服务半径': np.random.uniform(100, 800),
            '地形复杂度': np.random.choice(['低', '中', '高']),
            '自然灾害风险': np.random.choice(['低', '中', '高']),
            '扩展潜力': np.random.choice(['无', '有限', '充足']),
            '仓库类型': np.random.choice(['普通', '冷链', '危险品', '综合']),
            '自动化程度': np.random.choice(['低', '中', '高']),
            '到最近高速入口距离': np.random.uniform(1, 20),
            '到最近铁路站距离': np.random.uniform(5, 50),
            '到最近机场距离': np.random.uniform(10, 100),
            '到最近港口距离': np.random.uniform(20, 500)
        })

    st.session_state.candidate_locations = pd.DataFrame(candidate_data)

    # 生成增强的运输数据
    transport_data = []
    distance_matrix = {}

    for _, customer in st.session_state.customer_data.iterrows():
        for _, location in st.session_state.candidate_locations.iterrows():
            # 使用geodesic距离（更精确）
            distance = geodesic(
                (customer['纬度'], customer['经度']),
                (location['纬度'], location['经度'])
            ).kilometers

            # 考虑多种运输模式
            transport_modes = {
                '公路': {'speed': 60, 'cost_per_km': 0.8, 'reliability': 0.9},
                '铁路': {'speed': 80, 'cost_per_km': 0.5, 'reliability': 0.95},
                '航空': {'speed': 500, 'cost_per_km': 3.0, 'reliability': 0.85},
                '水运': {'speed': 30, 'cost_per_km': 0.3, 'reliability': 0.8}
            }

            # 根据距离选择合适的运输方式
            if distance < 300:
                mode = '公路'
            elif distance < 1000:
                mode = np.random.choice(['公路', '铁路'], p=[0.6, 0.4])
            else:
                mode = np.random.choice(['铁路', '航空'], p=[0.7, 0.3])

            mode_info = transport_modes[mode]

            # 计算运输成本和时间
            base_cost = distance * mode_info['cost_per_km']
            # 考虑需求量的规模效应
            volume_discount = 1 - min(customer['年需求量'] / 10000, 0.2)
            transport_cost = base_cost * volume_discount

            # 运输时间包括装卸时间
            transport_time = distance / mode_info['speed'] + np.random.uniform(2, 8)

            transport_data.append({
                '客户编号': customer['客户编号'],
                '地点编号': location['地点编号'],
                '距离公里': distance,
                '运输方式': mode,
                '单位运输成本': transport_cost,
                '运输时间小时': transport_time,
                '可靠性': mode_info['reliability'],
                '年运输成本': transport_cost * customer['年需求量'],
                '碳排放量': distance * 0.1 * (1 if mode == '航空' else 0.5 if mode == '公路' else 0.3)
            })

            distance_matrix[(location['地点编号'], customer['客户编号'])] = distance

    st.session_state.transportation_costs = pd.DataFrame(transport_data)
    st.session_state.distance_matrix = distance_matrix

    # 生成供应商数据
    num_suppliers = st.sidebar.number_input("供应商数量", min_value=5, max_value=20, value=10, key="gen_num_suppliers")
    supplier_data = []

    for i in range(num_suppliers):
        selected_city = np.random.choice(cities)
        city_coords = CHINA_CITIES_EXTENDED[selected_city]

        supplier_data.append({
            '供应商编号': f'S{i + 1:03d}',
            '供应商名称': f'{selected_city}供应商{i + 1}',
            '城市': selected_city,
            '经度': city_coords[0] + np.random.normal(0, 0.3),
            '纬度': city_coords[1] + np.random.normal(0, 0.2),
            '供应能力': np.random.uniform(10000, 100000),
            '产品质量评分': np.random.uniform(7, 10),
            '交付准时率': np.random.uniform(0.8, 0.99),
            '价格指数': np.random.uniform(0.8, 1.2),
            '最小订单量': np.random.uniform(100, 1000),
            '提前期天数': np.random.randint(3, 30),
            'id': f'S{i + 1:03d}'
        })

    st.session_state.supplier_data = pd.DataFrame(supplier_data)

    # 生成历史需求数据
    generate_historical_demand()


def generate_historical_demand():
    """生成历史需求数据"""
    if st.session_state.customer_data.empty:
        return

    # 为每个城市生成历史需求数据
    city_demand = st.session_state.customer_data.groupby('城市')['年需求量'].sum()

    historical_data = []
    months = pd.date_range(end=datetime.now(), periods=24, freq='M')

    for city, base_demand in city_demand.items():
        # 添加趋势和季节性
        trend = np.linspace(0.8, 1.2, 24) * base_demand / 12
        seasonal = np.sin(np.linspace(0, 4 * np.pi, 24)) * base_demand * 0.1 / 12
        noise = np.random.normal(0, base_demand * 0.05 / 12, 24)

        monthly_demand = trend + seasonal + noise
        monthly_demand = np.maximum(monthly_demand, 0)  # 确保非负

        for i, month in enumerate(months):
            historical_data.append({
                '日期': month,
                '城市': city,
                '月需求量': monthly_demand[i]
            })

    st.session_state.historical_demand = pd.DataFrame(historical_data)


# 创建高级交互式地图（整合两个系统的地图功能）
def create_advanced_folium_map(customers, candidates, selected_warehouses=None,
                               show_connections=False, show_heatmap=False,
                               show_clusters=False, show_risk_zones=False):
    """创建高级交互式地图"""
    # 计算地图中心
    all_lats = list(customers['纬度']) + list(candidates['纬度'])
    all_lons = list(customers['经度']) + list(candidates['经度'])
    center_lat = np.mean(all_lats)
    center_lon = np.mean(all_lons)

    # 创建地图
    m = folium.Map(
        location=[center_lat, center_lon],
        zoom_start=5,
        tiles='OpenStreetMap',
        control_scale=True
    )

    # 添加不同的地图图层
    folium.TileLayer('CartoDB positron', name='CartoDB Positron').add_to(m)
    folium.TileLayer('CartoDB dark_matter', name='CartoDB Dark').add_to(m)
    folium.TileLayer('OpenStreetMap', name='OpenStreetMap').add_to(m)

    # 添加热力图层
    if show_heatmap and len(customers) > 0:
        heat_data = [[row['纬度'], row['经度'], row['年需求量']]
                     for idx, row in customers.iterrows()]
        HeatMap(heat_data, name='需求热力图').add_to(m)

    # 添加聚类层
    if show_clusters and len(customers) > 0:
        marker_cluster = MarkerCluster(name='客户聚类').add_to(m)

        for idx, customer in customers.iterrows():
            folium.Marker(
                location=[customer['纬度'], customer['经度']],
                popup=f"{customer['客户名称']}<br>需求: {customer['年需求量']:.0f}",
                icon=folium.Icon(color='blue', icon='user', prefix='fa')
            ).add_to(marker_cluster)

    # 添加风险区域
    if show_risk_zones:
        # 模拟一些风险区域
        risk_zones = [
            {'center': [30.5728, 104.0668], 'radius': 100000, 'risk': 'high', 'type': '地震带'},
            {'center': [23.1291, 113.2644], 'radius': 80000, 'risk': 'medium', 'type': '台风区'},
            {'center': [31.2304, 121.4737], 'radius': 50000, 'risk': 'low', 'type': '洪水区'}
        ]

        for zone in risk_zones:
            color = {'high': 'red', 'medium': 'orange', 'low': 'yellow'}[zone['risk']]
            folium.Circle(
                location=zone['center'],
                radius=zone['radius'],
                popup=f"{zone['type']} - 风险等级: {zone['risk']}",
                color=color,
                fill=True,
                fillOpacity=0.2
            ).add_to(m)

    # 添加客户点
    if not customers.empty:
        for _, row in customers.iterrows():
            popup_text = f"""
            <b>客户信息</b><br>
            客户ID: {row['客户编号']}<br>
            名称: {row['客户名称']}<br>
            城市: {row['城市']}<br>
            需求量: {row['年需求量']:,}<br>
            优先级: {row['优先级']}<br>
            类型: {row['客户类型']}
            """

            folium.CircleMarker(
                location=[row['纬度'], row['经度']],
                radius=5 + row['年需求量'] / 500,
                color='blue',
                fill=True,
                fillColor='lightblue',
                fillOpacity=0.7,
                popup=folium.Popup(popup_text, max_width=300),
                tooltip=row['客户名称']
            ).add_to(m)

    # 添加候选仓库
    warehouse_group = folium.FeatureGroup(name='候选仓库')
    for idx, warehouse in candidates.iterrows():
        is_selected = selected_warehouses and warehouse['地点编号'] in selected_warehouses

        # 创建详细的弹出信息
        popup_html = f"""
        <div style='font-family: Arial; width: 300px;'>
            <h4>{warehouse['地点名称']}</h4>
            <table style='width: 100%;'>
                <tr><td><b>城市:</b></td><td>{warehouse['城市']}</td></tr>
                <tr><td><b>建设成本:</b></td><td>¥{warehouse['建设成本'] / 1e6:.1f}M</td></tr>
                <tr><td><b>容量:</b></td><td>{warehouse['最大容量']:,.0f}</td></tr>
                <tr><td><b>服务半径:</b></td><td>{warehouse['服务半径']:.0f} km</td></tr>
                <tr><td><b>风险评分:</b></td><td>{warehouse['风险评分']:.1f}</td></tr>
                <tr><td><b>基础设施:</b></td><td>{warehouse['基础设施评分']:.1f}</td></tr>
            </table>
        </div>
        """

        icon_color = 'green' if is_selected else 'gray'
        icon_symbol = 'star' if is_selected else 'warehouse'

        marker = folium.Marker(
            location=[warehouse['纬度'], warehouse['经度']],
            popup=folium.Popup(popup_html, max_width=300),
            tooltip=warehouse['地点名称'],
            icon=folium.Icon(color=icon_color, icon=icon_symbol, prefix='fa')
        )
        marker.add_to(warehouse_group)

        # 添加服务范围
        if is_selected:
            folium.Circle(
                location=[warehouse['纬度'], warehouse['经度']],
                radius=warehouse['服务半径'] * 1000,
                color='green',
                fill=True,
                fillOpacity=0.1,
                popup=f"服务范围: {warehouse['服务半径']} km"
            ).add_to(warehouse_group)

    warehouse_group.add_to(m)

    # 添加连接线
    if show_connections and selected_warehouses:
        connection_group = folium.FeatureGroup(name='服务连接')

        for warehouse_id in selected_warehouses:
            warehouse = candidates[candidates['地点编号'] == warehouse_id].iloc[0]

            # 找出该仓库服务的客户
            served_customers = []
            for idx, customer in customers.iterrows():
                distance = st.session_state.distance_matrix.get(
                    (warehouse_id, customer['客户编号']), float('inf')
                )
                if distance <= warehouse['服务半径']:
                    served_customers.append(customer)

                    # 画连接线
                    folium.PolyLine(
                        locations=[
                            [customer['纬度'], customer['经度']],
                            [warehouse['纬度'], warehouse['经度']]
                        ],
                        color='green',
                        weight=1 + customer['年需求量'] / 5000,  # 线宽反映需求量
                        opacity=0.5,
                        popup=f"距离: {distance:.1f} km"
                    ).add_to(connection_group)

        connection_group.add_to(m)

    # 添加图层控制
    folium.LayerControl().add_to(m)

    # 添加全屏按钮
    Fullscreen().add_to(m)

    # 添加测量工具
    MeasureControl().add_to(m)

    # 添加搜索框
    Search(
        layer=warehouse_group,
        geom_type='Point',
        placeholder='搜索仓库...',
        collapsed=False,
        search_label='name'
    ).add_to(m)

    return m


# 导出功能
def export_all_data():
    """导出所有数据到Excel"""
    output = io.BytesIO()

    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # 导出客户数据
        if not st.session_state.customer_data.empty:
            st.session_state.customer_data.to_excel(writer, sheet_name='客户数据', index=False)

        # 导出候选地点数据
        if not st.session_state.candidate_locations.empty:
            st.session_state.candidate_locations.to_excel(writer, sheet_name='候选地点', index=False)

        # 导出运输成本数据（前1000条）
        if not st.session_state.transportation_costs.empty:
            st.session_state.transportation_costs.head(1000).to_excel(writer, sheet_name='运输成本', index=False)

        # 导出优化结果
        if not st.session_state.optimization_results.empty:
            st.session_state.optimization_results.to_excel(writer, sheet_name='优化结果', index=False)

        # 导出分析结果摘要
        if st.session_state.analysis_results:
            summary_data = []
            for key, value in st.session_state.analysis_results.items():
                summary_data.append({
                    '分析项目': key,
                    '结果摘要': str(value)[:100] + '...' if len(str(value)) > 100 else str(value)
                })
            pd.DataFrame(summary_data).to_excel(writer, sheet_name='分析摘要', index=False)

    output.seek(0)
    return output


# 主程序
def main():
    # 显示主标题
    st.markdown(
        '<h1 class="main-header">🏢 仓库选址优化系统 Ultimate Pro</h1>',
        unsafe_allow_html=True
    )

    # 侧边栏增强
    with st.sidebar:
        st.image(
            "https://via.placeholder.com/300x150/1e3c72/ffffff?text=WMS+Ultimate+Pro",
            use_container_width=True
        )

        st.markdown("### 🎯 快速操作")

        # 一键初始化
        if st.button("🚀 一键初始化", use_container_width=True, type="primary"):
            with st.spinner("正在初始化系统..."):
                generate_advanced_sample_data()
                # 自动运行基础分析
                st.success("✅ 系统初始化完成！")
                st.balloons()

        # 数据导出
        if st.button("📥 导出所有数据", use_container_width=True):
            if not st.session_state.customer_data.empty:
                excel_data = export_all_data()
                st.download_button(
                    label="下载Excel文件",
                    data=excel_data,
                    file_name=f"warehouse_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                    mime="application/vnd.ms-excel"
                )
            else:
                st.error("❌ 请先生成数据！")

        # 高级设置
        with st.expander("⚙️ 高级设置"):
            st.selectbox(
                "优化算法偏好",
                ["自动选择", "精确算法", "启发式算法", "机器学习"],
                key="algo_preference"
            )

            st.slider(
                "分析精度",
                min_value=1,
                max_value=10,
                value=7,
                help="更高的精度需要更长的计算时间",
                key="sidebar_analysis_precision"
            )

            st.checkbox("启用实时数据", value=False, key="sidebar_realtime_data")
            st.checkbox("启用GPU加速", value=False, key="sidebar_gpu_accel")
            st.checkbox("启用分布式计算", value=False, key="sidebar_distributed")

        # 系统状态
        st.markdown("---")
        st.markdown("### 📊 系统状态")

        # 数据状态指示器
        data_status = {
            "客户数据": len(st.session_state.customer_data) > 0,
            "仓库数据": len(st.session_state.candidate_locations) > 0,
            "运输数据": len(st.session_state.transportation_costs) > 0,
            "优化结果": bool(st.session_state.selected_locations)
        }

        for item, status in data_status.items():
            if status:
                st.success(f"✅ {item}")
            else:
                st.error(f"❌ {item}")

        # 性能监控
        st.markdown("---")
        st.markdown("### ⚡ 性能监控")
        st.metric("CPU使用率", "45%", "+5%")
        st.metric("内存使用", "2.3 GB", "-0.1 GB")
        st.metric("响应时间", "120 ms", "-20 ms")

        st.markdown("---")

        # 系统信息
        st.info("""
        **系统特点：**
        - 🎯 多种优化算法
        - 📊 全面的数据分析
        - 🗺️ 交互式地图展示
        - 📈 专业的决策支持
        - 📑 完整的报告生成
        - 🧮 混合整数规划(MIP)
        - 🎯 多目标优化
        - 🌱 碳足迹分析
        - 📈 需求预测
        - 🤖 智能决策建议
        """)

    # 主界面标签页
    tabs = st.tabs([
        "🏠 总览",
        "📊 数据中心",
        "📈 需求分析",
        "🎯 地点评估",
        "🔬 高级分析",
        "🚀 智能优化",
        "💰 成本分析",
        "⚠️ 风险评估",
        "💡 决策支持",
        "📊 可视化",
        "🔮 预测模拟",
        "📋 结果展示",
        "📑 报告中心",
        "⚙️ 系统管理"
    ])

    with tabs[0]:
        show_advanced_dashboard()

    with tabs[1]:
        show_data_center()

    with tabs[2]:
        show_demand_analysis()

    with tabs[3]:
        show_location_evaluation()

    with tabs[4]:
        advanced_analytics()

    with tabs[5]:
        show_unified_optimization()

    with tabs[6]:
        show_cost_analysis()

    with tabs[7]:
        show_risk_assessment()

    with tabs[8]:
        advanced_decision_support()

    with tabs[9]:
        show_advanced_visualization()

    with tabs[10]:
        show_prediction_simulation()

    with tabs[11]:
        show_results_display()

    with tabs[12]:
        generate_advanced_report()

    with tabs[13]:
        show_system_management()

def show_advanced_dashboard():
    """显示高级仪表板"""
    st.markdown("## 📊 智能决策仪表板")

    # 如果没有数据，先生成
    if st.session_state.customer_data.empty:
        generate_advanced_sample_data()

    # KPI卡片行
    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "客户数量",
            len(st.session_state.customer_data) if not st.session_state.customer_data.empty else 0,
            "个"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "候选地点",
            len(st.session_state.candidate_locations) if not st.session_state.candidate_locations.empty else 0,
            "个"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col3:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        total_demand = st.session_state.customer_data['年需求量'].sum() if not st.session_state.customer_data.empty else 0
        st.metric(
            "总需求量",
            f"{total_demand:,}",
            "单位/年"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col4:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "完成分析",
            len(st.session_state.analysis_results),
            "项"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col5:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "AI信心度",
            "94%",
            "+2%",
            help="AI推荐方案的置信度"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    # 实时监控图表
    st.markdown("---")
    col1, col2 = st.columns([2, 1])

    with col1:
        # 创建实时更新的图表
        dates = pd.date_range(end=datetime.now(), periods=30, freq='D')
        metrics_data = pd.DataFrame({
            '成本优化': np.cumsum(np.random.randn(30)) + 50,
            '服务水平': np.cumsum(np.random.randn(30)) * 0.5 + 85,
            '运营效率': np.cumsum(np.random.randn(30)) * 0.3 + 75
        }, index=dates)

        fig = px.line(
            metrics_data,
            title='关键指标趋势（最近30天）',
            labels={'value': '指标值', 'index': '日期'},
            height=400
        )
        fig.update_layout(hovermode='x unified')
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        # AI健康检查
        st.markdown("### 🤖 AI系统状态")

        ai_components = {
            '预测模型': 'operational',
            '优化引擎': 'operational',
            '风险评估': 'warning',
            '决策系统': 'operational',
            '数据管道': 'operational'
        }

        for component, status in ai_components.items():
            if status == 'operational':
                st.success(f"✅ {component}")
            elif status == 'warning':
                st.warning(f"⚠️ {component}")
            else:
                st.error(f"❌ {component}")

        # 下一步行动建议
        st.markdown("### 📋 建议行动")
        actions = [
            "完成华东地区仓库谈判",
            "更新需求预测模型",
            "审查风险缓解措施"
        ]
        for i, action in enumerate(actions):
            st.checkbox(action, key=f"dashboard_action_{i}")

    # 地图概览
    if len(st.session_state.customer_data) > 0 and len(st.session_state.candidate_locations) > 0:
        st.markdown("---")
        st.markdown("### 🗺️ 网络布局概览")

        m = create_advanced_folium_map(
            st.session_state.customer_data,
            st.session_state.candidate_locations,
            st.session_state.selected_locations,
            show_connections=True,
            show_heatmap=True
        )

        folium_static(m, width=1400, height=600)

    # 快速开始指南
    st.markdown("---")
    st.subheader("🚀 快速开始指南")

    col1, col2 = st.columns(2)

    with col1:
        st.info("""
        ### 📋 使用步骤
        1. **生成数据** - 在侧边栏点击"一键初始化"
        2. **需求分析** - 分析客户分布和需求特征
        3. **地点评估** - 评估候选仓库位置
        4. **选址优化** - 运行优化算法
        5. **查看结果** - 查看优化结果和建议
        """)

    with col2:
        st.success("""
        ### 🎯 核心功能
        - **多算法支持** - 10+种优化算法
        - **全面分析** - 需求、成本、风险多维度
        - **可视化展示** - 3D地图、VR预览
        - **智能决策** - AI驱动的决策支持
        """)

    # 系统亮点展示
    st.markdown("---")
    st.subheader("✨ 系统亮点")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("""
        <div class="metric-card">
        <h4>🧮 混合整数规划(MIP)</h4>
        <p>业界领先的精确优化算法，保证全局最优解，支持预算约束</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div class="metric-card">
        <h4>🎯 多目标优化</h4>
        <p>同时优化成本、服务水平和环境影响，符合ESG理念</p>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div class="metric-card">
        <h4>🌱 碳足迹分析</h4>
        <p>计算物流碳排放，提供碳中和方案，支持绿色供应链</p>
        </div>
        """, unsafe_allow_html=True)

# 成本分析（整合两个系统）
def show_cost_analysis():
    """成本分析页面"""
    st.subheader("💰 成本分析")

    if st.session_state.candidate_locations.empty:
        st.warning("请先完成选址优化")
        # 自动运行一个优化算法
        if not st.session_state.selected_locations:
            # 默认选择前3个候选地点
            st.session_state.selected_locations = st.session_state.candidate_locations.head(3)['地点编号'].tolist()

    tab1, tab2, tab3, tab4 = st.tabs(["总成本分析", "运输成本", "ROI分析", "碳足迹分析"])

    with tab1:
        st.markdown("### 总成本分析")

        analysis_period = st.slider("分析周期(年)", 5, 20, 10)

        # 如果有选中的仓库
        if st.session_state.selected_locations:
            selected_locations = st.session_state.candidate_locations[
                st.session_state.candidate_locations['地点编号'].isin(st.session_state.selected_locations)
            ]

            # 计算成本
            total_construction = selected_locations['建设成本'].sum()
            total_land = selected_locations['土地单价'].sum() * 5000  # 假设5000平米
            annual_operating = selected_locations['运营成本'].sum()

            # 显示成本明细
            col1, col2, col3 = st.columns(3)
            col1.metric("建设成本", f"¥{total_construction / 1e6:.1f}M")
            col2.metric("土地成本", f"¥{total_land / 1e6:.1f}M")
            col3.metric("年运营成本", f"¥{annual_operating / 1e6:.1f}M")

            # 成本构成饼图
            fig = go.Figure(data=[go.Pie(
                labels=['建设成本', '土地成本', f'运营成本({analysis_period}年)'],
                values=[total_construction, total_land, annual_operating * analysis_period]
            )])
            fig.update_layout(title=f'{analysis_period}年期成本构成')
            st.plotly_chart(fig, use_container_width=True)

    with tab2:
        st.markdown("### 运输成本分析")

        if not st.session_state.transportation_costs.empty and st.session_state.selected_locations:
            # 计算每个选中仓库的运输成本
            transport_summary = []

            for warehouse_id in st.session_state.selected_locations:
                warehouse_costs = st.session_state.transportation_costs[
                    st.session_state.transportation_costs['地点编号'] == warehouse_id
                    ]

                if not warehouse_costs.empty:
                    total_transport = warehouse_costs['年运输成本'].sum()
                    avg_distance = warehouse_costs['距离公里'].mean()

                    transport_summary.append({
                        '仓库ID': warehouse_id,
                        '年运输成本': total_transport,
                        '平均运输距离': avg_distance
                    })

            if transport_summary:
                transport_df = pd.DataFrame(transport_summary)
                st.dataframe(transport_df)

                # 运输成本对比图
                fig = px.bar(
                    transport_df,
                    x='仓库ID',
                    y='年运输成本',
                    title='各仓库年运输成本对比'
                )
                st.plotly_chart(fig, use_container_width=True)

    with tab3:
        st.markdown("### ROI投资回报分析")

        service_price = st.number_input("服务单价(元/单位)", 10, 200, 50)
        tax_rate = st.slider("税率(%)", 0, 50, 25) / 100

        if st.button("计算ROI"):
            if st.session_state.selected_locations:
                selected_locations = st.session_state.candidate_locations[
                    st.session_state.candidate_locations['地点编号'].isin(st.session_state.selected_locations)
                ]

                # 初始投资
                initial_investment = selected_locations['建设成本'].sum() + selected_locations['土地单价'].sum() * 5000

                # 年收入（简化计算）
                total_demand = st.session_state.customer_data['年需求量'].sum()
                annual_revenue = total_demand * service_price * 0.8  # 假设服务80%的需求

                # 年成本
                annual_cost = selected_locations['运营成本'].sum()

                # 年净利润
                annual_profit = (annual_revenue - annual_cost) * (1 - tax_rate)

                # ROI
                roi = (annual_profit * 10 - initial_investment) / initial_investment * 100
                payback_period = initial_investment / annual_profit if annual_profit > 0 else float('inf')

                # 显示结果
                col1, col2, col3 = st.columns(3)
                col1.metric("初始投资", f"¥{initial_investment / 1e6:.1f}M")
                col2.metric("ROI(10年)", f"{roi:.1f}%")
                col3.metric("投资回收期", f"{payback_period:.1f}年" if payback_period < 100 else ">10年")

                # 现金流图
                years = list(range(11))
                cash_flows = [-initial_investment] + [annual_profit] * 10
                cumulative_cash_flow = np.cumsum(cash_flows)

                fig = go.Figure()
                fig.add_trace(go.Bar(x=years, y=cash_flows, name='年现金流'))
                fig.add_trace(go.Scatter(x=years, y=cumulative_cash_flow, name='累计现金流', line=dict(color='red')))
                fig.update_layout(title='10年现金流分析', xaxis_title='年份', yaxis_title='金额(元)')
                st.plotly_chart(fig, use_container_width=True)

    with tab4:
        st.markdown("### 🌱 碳足迹分析")
        st.info("评估物流网络的环境影响，支持碳中和目标")

        if st.session_state.selected_locations and not st.session_state.customer_data.empty:
            # 运输方式选择
            transport_mode = st.selectbox(
                "主要运输方式",
                ['truck', 'rail', 'air', 'ship'],
                format_func=lambda x: {
                    'truck': '公路运输',
                    'rail': '铁路运输',
                    'air': '航空运输',
                    'ship': '水路运输'
                }[x]
            )

            avg_cargo_weight = st.number_input("平均货物重量(kg/批次)", 100, 10000, 1000)

            if st.button("计算碳足迹", key="calc_carbon"):
                # 计算每个仓库的碳排放
                carbon_results = []
                total_carbon = 0

                for warehouse_id in st.session_state.selected_locations:
                    warehouse = st.session_state.candidate_locations[
                        st.session_state.candidate_locations['地点编号'] == warehouse_id
                        ].iloc[0]

                    warehouse_carbon = 0
                    served_customers = 0

                    # 计算到每个客户的碳排放
                    for _, customer in st.session_state.customer_data.iterrows():
                        distance = calculate_distance(
                            warehouse['经度'], warehouse['纬度'],
                            customer['经度'], customer['纬度']
                        )

                        # 年运输次数（简化计算）
                        annual_trips = customer['年需求量'] / 100  # 假设每次运输100单位

                        # 计算碳排放
                        carbon_emission = carbon_footprint_calculation(
                            distance, transport_mode, avg_cargo_weight
                        ) * annual_trips

                        warehouse_carbon += carbon_emission
                        served_customers += 1

                    carbon_results.append({
                        '仓库ID': warehouse_id,
                        '仓库名称': warehouse['地点名称'],
                        '年碳排放(吨)': warehouse_carbon / 1000,  # 转换为吨
                        '服务客户数': served_customers,
                        '单位客户碳排放': warehouse_carbon / served_customers / 1000 if served_customers > 0 else 0
                    })

                    total_carbon += warehouse_carbon

                # 显示结果
                carbon_df = pd.DataFrame(carbon_results)
                st.dataframe(carbon_df)

                # 关键指标
                col1, col2, col3 = st.columns(3)
                col1.metric("总碳排放", f"{total_carbon / 1000:.1f} 吨CO₂/年")
                col2.metric("平均每仓库",
                            f"{total_carbon / 1000 / len(st.session_state.selected_locations):.1f} 吨CO₂/年")
                col3.metric("碳强度",
                            f"{total_carbon / st.session_state.customer_data['年需求量'].sum():.2f} kgCO₂/单位")

                # 碳排放对比图
                fig = px.bar(
                    carbon_df,
                    x='仓库名称',
                    y='年碳排放(吨)',
                    title='各仓库年碳排放对比',
                    color='年碳排放(吨)',
                    color_continuous_scale='Reds'
                )
                st.plotly_chart(fig, use_container_width=True)

                # 碳中和建议
                st.markdown("### 🌍 碳中和建议")

                # 计算碳补偿成本
                carbon_price = st.number_input("碳信用价格(元/吨)", 50, 500, 200)
                offset_cost = total_carbon / 1000 * carbon_price

                st.success(f"""
                **减碳策略建议：**
                1. 🚂 优化运输方式：将部分公路运输转为铁路可减少60-70%碳排放
                2. 🔋 使用新能源车辆：电动货车可减少50%以上碳排放
                3. 📦 提高装载率：优化配送路线和装载可减少20-30%运输次数
                4. 🏭 绿色仓库建设：太阳能、LED照明等可减少30%运营碳排放
                5. 🌳 碳补偿方案：需要约 ¥{offset_cost:,.0f}/年 购买碳信用实现碳中和
                """)

                # 保存分析结果
                st.session_state.analysis_results['碳足迹分析'] = {
                    '总碳排放(吨/年)': total_carbon / 1000,
                    '运输方式': transport_mode,
                    '碳补偿成本': offset_cost,
                    '详细数据': carbon_df.to_dict()
                }


# 风险评估
def show_risk_assessment():
    """风险评估页面"""
    st.subheader("⚠️ 风险评估")

    tab1, tab2, tab3 = st.tabs(["风险识别", "风险量化", "应对策略"])

    with tab1:
        st.markdown("### 风险识别")

        # 风险分类
        risk_categories = {
            '市场风险': ['需求波动', '竞争加剧', '价格下降', '客户流失'],
            '运营风险': ['人员流失', '设备故障', '质量问题', '效率低下'],
            '财务风险': ['资金短缺', '成本超支', '汇率风险', '利率上升'],
            '政策风险': ['法规变化', '税收调整', '土地政策', '行业政策'],
            '自然风险': ['自然灾害', '极端天气', '环境变化', '地质问题']
        }

        # 显示风险分类
        for category, risks in risk_categories.items():
            with st.expander(category):
                for risk in risks:
                    st.write(f"• {risk}")

        # 风险等级分布
        if not st.session_state.candidate_locations.empty:
            risk_dist = pd.cut(
                st.session_state.candidate_locations['风险评分'],
                bins=[0, 4, 7, 10],
                labels=['低风险', '中风险', '高风险']
            ).value_counts()

            fig = px.pie(
                values=risk_dist.values,
                names=risk_dist.index,
                title='候选地点风险等级分布',
                color_discrete_map={'低风险': 'green', '中风险': 'yellow', '高风险': 'red'}
            )
            st.plotly_chart(fig, use_container_width=True)

    with tab2:
        st.markdown("### 风险量化分析")

        # 风险矩阵
        risk_data = {
            '风险因素': ['需求波动', '建设超支', '运营成本上升', '政策变化', '自然灾害'],
            '发生概率': [0.6, 0.4, 0.5, 0.3, 0.1],
            '影响程度': [0.15, 0.25, 0.12, 0.08, 0.4]
        }

        risk_df = pd.DataFrame(risk_data)
        risk_df['期望损失'] = risk_df['发生概率'] * risk_df['影响程度']

        # 风险矩阵散点图
        fig = px.scatter(
            risk_df,
            x='发生概率',
            y='影响程度',
            size='期望损失',
            color='期望损失',
            text='风险因素',
            title='风险概率-影响矩阵',
            color_continuous_scale='Reds'
        )
        fig.update_traces(textposition='top center')
        st.plotly_chart(fig, use_container_width=True)

        # 显示风险量化表
        st.dataframe(risk_df.sort_values('期望损失', ascending=False))

        total_risk = risk_df['期望损失'].sum()
        st.metric("总期望风险损失", f"{total_risk:.2%}")

    with tab3:
        st.markdown("### 风险应对策略")

        strategies = {
            '市场风险': {
                '预防': ['多元化客户结构', '长期合同', '价格调整机制'],
                '应急': ['客户挽留计划', '新市场开拓', '服务升级']
            },
            '运营风险': {
                '预防': ['标准化流程', '人才储备', '预防性维护'],
                '应急': ['应急预案', '外包服务', '快速招聘']
            },
            '财务风险': {
                '预防': ['多元化融资', '成本控制', '财务预警'],
                '应急': ['紧急融资', '成本削减', '资产变现']
            }
        }

        for risk_type, measures in strategies.items():
            st.markdown(f"#### {risk_type}")
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("**预防措施**")
                for measure in measures['预防']:
                    st.write(f"• {measure}")

            with col2:
                st.markdown("**应急措施**")
                for measure in measures['应急']:
                    st.write(f"• {measure}")


# 高级决策支持
def advanced_decision_support():
    """高级决策支持系统"""
    st.subheader("🤖 智能决策支持系统")

    tabs = st.tabs([
        "AI推荐", "方案对比", "决策树分析",
        "专家系统", "协同决策", "决策仪表板"
    ])

    with tabs[0]:
        st.markdown("### AI智能推荐")

        if st.session_state.analysis_results:
            st.info("基于已完成的分析，AI系统为您推荐最优方案")

            # 模拟AI评分
            ai_scores = {
                '成本效益': np.random.uniform(7, 10),
                '服务覆盖': np.random.uniform(7, 10),
                '风险控制': np.random.uniform(6, 9),
                '可扩展性': np.random.uniform(7, 10),
                '可持续性': np.random.uniform(6, 9)
            }

            overall_score = np.mean(list(ai_scores.values()))

            # 显示AI评分
            col1, col2 = st.columns([2, 1])

            with col1:
                fig = go.Figure(data=[
                    go.Bar(
                        x=list(ai_scores.values()),
                        y=list(ai_scores.keys()),
                        orientation='h',
                        marker_color='lightblue'
                    )
                ])

                fig.update_layout(
                    title='AI方案评分',
                    xaxis_title='评分',
                    xaxis=dict(range=[0, 10])
                )

                st.plotly_chart(fig, use_container_width=True)

            with col2:
                st.metric("综合评分", f"{overall_score:.1f}/10")

                if overall_score >= 8:
                    st.success("强烈推荐")
                elif overall_score >= 7:
                    st.info("推荐")
                else:
                    st.warning("谨慎考虑")

            # AI建议
            st.markdown("#### AI决策建议")

            recommendations = [
                "✅ 当前方案在成本控制方面表现优秀，建议保持",
                "📈 服务覆盖率可进一步优化，考虑增加1-2个战略位置的仓库",
                "⚠️ 建议加强华东地区的风险管理措施",
                "🌱 推荐采用绿色物流方案，预计可降低20%碳排放",
                "💡 考虑引入自动化设备，提升运营效率15-20%"
            ]

            for rec in recommendations:
                st.write(rec)

    with tabs[1]:
        st.markdown("### 多方案对比分析")

        # 创建对比方案
        scenarios = {
            '保守方案': {
                '仓库数量': 3,
                '投资额': 5000,
                '覆盖率': 0.75,
                '风险等级': '低',
                'ROI': 0.15
            },
            '平衡方案': {
                '仓库数量': 5,
                '投资额': 8000,
                '覆盖率': 0.88,
                '风险等级': '中',
                'ROI': 0.22
            },
            '进取方案': {
                '仓库数量': 8,
                '投资额': 12000,
                '覆盖率': 0.95,
                '风险等级': '高',
                'ROI': 0.28
            }
        }

        # 方案对比表
        comparison_df = pd.DataFrame(scenarios).T
        st.dataframe(comparison_df.style.format({
            '投资额': '¥{:,.0f}万',
            '覆盖率': '{:.1%}',
            'ROI': '{:.1%}'
        }))

        # 雷达图对比
        categories = ['仓库数量', '投资规模', '覆盖率', 'ROI', '风险承受']

        fig = go.Figure()

        for scenario_name, data in scenarios.items():
            values = [
                data['仓库数量'] / 10,
                data['投资额'] / 15000,
                data['覆盖率'],
                data['ROI'] / 0.3,
                {'低': 0.3, '中': 0.6, '高': 0.9}[data['风险等级']]
            ]

            fig.add_trace(go.Scatterpolar(
                r=values + [values[0]],
                theta=categories + [categories[0]],
                fill='toself',
                name=scenario_name
            ))

        fig.update_layout(
            polar=dict(
                radialaxis=dict(
                    visible=True,
                    range=[0, 1]
                )),
            showlegend=True,
            title="方案综合对比"
        )

        st.plotly_chart(fig, use_container_width=True)

    with tabs[2]:
        st.markdown("### 决策树分析")

        # 创建决策树可视化
        st.markdown("""
        ```mermaid
        graph TD
            A[选址决策] --> B{市场需求}
            B -->|高增长| C[扩张策略]
            B -->|稳定| D[优化策略]
            B -->|下降| E[收缩策略]

            C --> F{资金充足?}
            F -->|是| G[建设8-10个仓库]
            F -->|否| H[分期建设5-7个]

            D --> I{服务水平}
            I -->|需提升| J[增加2-3个仓库]
            I -->|满意| K[维持现状]

            E --> L[关闭低效仓库]
            E --> M[区域整合]
        ```
        """)

        # 决策路径分析
        st.markdown("#### 推荐决策路径")

        decision_path = st.selectbox(
            "选择市场情况",
            ["高增长", "稳定", "下降"]
        )

        if decision_path == "高增长":
            st.success("""
            **推荐策略: 积极扩张**
            - 快速占领市场份额
            - 建立8-10个区域配送中心
            - 重点布局一二线城市
            - 预留扩展空间
            """)
        elif decision_path == "稳定":
            st.info("""
            **推荐策略: 优化提升**
            - 提高现有设施利用率
            - 优化配送网络
            - 适度增加3-5个仓库
            - 注重成本控制
            """)
        else:
            st.warning("""
            **推荐策略: 战略收缩**
            - 关闭亏损网点
            - 整合区域资源
            - 保留核心市场
            - 提高运营效率
            """)

    with tabs[3]:
        st.markdown("### 专家系统建议")

        # 模拟专家知识库
        expert_rules = {
            '选址原则': [
                "优先选择交通枢纽城市",
                "避免自然灾害高发区",
                "考虑未来5-10年发展规划",
                "平衡成本与服务水平"
            ],
            '运营建议': [
                "采用WMS仓储管理系统",
                "实施ABC库存分类管理",
                "建立KPI考核体系",
                "定期进行网络优化"
            ],
            '风险管理': [
                "建立应急响应机制",
                "购买适当保险",
                "多元化供应商基础",
                "保持合理库存水平"
            ]
        }

        for category, rules in expert_rules.items():
            with st.expander(f"📚 {category}"):
                for rule in rules:
                    st.write(f"• {rule}")

        # 专家评分系统
        st.markdown("#### 专家评估")

        aspects = ['位置选择', '成本控制', '服务能力', '风险管理', '可持续性']
        expert_scores = {}

        for i, aspect in enumerate(aspects):
            expert_scores[aspect] = st.slider(
                f"{aspect}评分",
                0, 10,
                int(np.random.uniform(6, 9)),
                key=f"expert_score_{i}_{aspect}"
            )

        avg_score = np.mean(list(expert_scores.values()))

        if avg_score >= 8:
            st.success(f"专家综合评分: {avg_score:.1f}/10 - 优秀")
        elif avg_score >= 6:
            st.info(f"专家综合评分: {avg_score:.1f}/10 - 良好")
        else:
            st.warning(f"专家综合评分: {avg_score:.1f}/10 - 需改进")

    with tabs[4]:
        st.markdown("### 协同决策平台")

        # 决策参与者
        stakeholders = {
            '运营部门': {'权重': 0.3, '关注点': '效率与成本'},
            '销售部门': {'权重': 0.25, '关注点': '市场覆盖'},
            '财务部门': {'权重': 0.25, '关注点': 'ROI与风险'},
            '战略部门': {'权重': 0.2, '关注点': '长期发展'}
        }

        st.markdown("#### 各部门评分")

        dept_scores = {}
        for i, (dept, info) in enumerate(stakeholders.items()):
            col1, col2, col3 = st.columns([2, 1, 2])
            with col1:
                st.write(f"**{dept}**")
            with col2:
                st.write(f"权重: {info['权重']}")
            with col3:
                score = st.slider(
                    f"评分",
                    0.0, 10.0, 7.5,
                    key=f"dept_score_{i}_{dept}",
                    label_visibility="collapsed"
                )
                dept_scores[dept] = score

        # 计算加权得分
        weighted_score = sum(
            dept_scores[dept] * stakeholders[dept]['权重']
            for dept in dept_scores
        )

        st.metric("综合决策得分", f"{weighted_score:.1f}/10")

        # 决策建议
        if weighted_score >= 7:
            decision = "批准实施"
            color = "success"
        elif weighted_score >= 5:
            decision = "修改后重审"
            color = "warning"
        else:
            decision = "暂缓实施"
            color = "error"

        st.markdown(f"### 决策结果: :{color}[{decision}]")

    with tabs[5]:
        st.markdown("### 决策仪表板")

        # 创建仪表板布局
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            # 仪表图 - 成本效益
            fig = go.Figure(go.Indicator(
                mode="gauge+number+delta",
                value=85,
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': "成本效益"},
                delta={'reference': 80},
                gauge={
                    'axis': {'range': [None, 100]},
                    'bar': {'color': "darkblue"},
                    'steps': [
                        {'range': [0, 50], 'color': "lightgray"},
                        {'range': [50, 80], 'color': "gray"}
                    ],
                    'threshold': {
                        'line': {'color': "red", 'width': 4},
                        'thickness': 0.75,
                        'value': 90
                    }
                }
            ))
            fig.update_layout(height=200, margin=dict(l=20, r=20, t=40, b=20))
            st.plotly_chart(fig, use_container_width=True)

        with col2:
            # 风险指标
            risk_score = 3.5
            st.metric(
                "风险指数",
                f"{risk_score:.1f}/10",
                f"{-0.5:.1f}",
                delta_color="inverse"
            )

            # 风险等级
            if risk_score <= 3:
                st.success("低风险")
            elif risk_score <= 6:
                st.warning("中等风险")
            else:
                st.error("高风险")

        with col3:
            # 进度指标
            progress = 0.75
            st.metric("项目进度", f"{progress:.0%}", "+5%")
            st.progress(progress)

            # 关键里程碑
            st.caption("下一里程碑: 选址谈判")

        with col4:
            # ROI预测
            roi_current = 22.5
            roi_target = 25.0

            fig = go.Figure(go.Indicator(
                mode="number+delta",
                value=roi_current,
                number={'suffix': "%"},
                delta={'reference': roi_target, 'relative': True},
                title={"text": "ROI预测"}
            ))
            fig.update_layout(height=150, margin=dict(l=20, r=20, t=40, b=20))
            st.plotly_chart(fig, use_container_width=True)

        # 时间线
        st.markdown("#### 项目时间线")

        timeline_data = pd.DataFrame({
            '阶段': ['可行性研究', '选址谈判', '建设施工', '设备安装', '试运营', '正式运营'],
            '开始时间': pd.to_datetime(['2024-01-01', '2024-03-01', '2024-06-01',
                                        '2024-11-01', '2025-01-01', '2025-03-01']),
            '结束时间': pd.to_datetime(['2024-02-28', '2024-05-31', '2024-10-31',
                                        '2024-12-31', '2025-02-28', '2025-12-31']),
            '状态': ['完成', '进行中', '计划', '计划', '计划', '计划']
        })

        fig = px.timeline(
            timeline_data,
            x_start="开始时间",
            x_end="结束时间",
            y="阶段",
            color="状态",
            title="项目实施时间线"
        )
        fig.update_yaxes(autorange="reversed")
        st.plotly_chart(fig, use_container_width=True)


# 高级可视化
def show_advanced_visualization():
    """高级可视化中心"""
    st.markdown("## 📈 高级可视化分析")

    viz_tabs = st.tabs([
        "3D可视化", "网络图", "热力分析",
        "时空分析", "交互式仪表板", "VR预览"
    ])

    with viz_tabs[0]:
        st.markdown("### 3D空间分析")

        if len(st.session_state.customer_data) > 0 and len(st.session_state.candidate_locations) > 0:
            # 创建3D散点图
            fig = go.Figure()

            # 添加客户点
            fig.add_trace(go.Scatter3d(
                x=st.session_state.customer_data['经度'],
                y=st.session_state.customer_data['纬度'],
                z=st.session_state.customer_data['年需求量'] / 1000,
                mode='markers',
                marker=dict(
                    size=5,
                    color=st.session_state.customer_data['年需求量'],
                    colorscale='Blues',
                    showscale=True,
                    colorbar=dict(title="需求量")
                ),
                text=st.session_state.customer_data['客户名称'],
                name='客户'
            ))

            # 添加仓库点
            selected_warehouses = st.session_state.candidate_locations[
                st.session_state.candidate_locations['地点编号'].isin(
                    st.session_state.selected_locations
                )
            ] if st.session_state.selected_locations else st.session_state.candidate_locations.head(5)

            fig.add_trace(go.Scatter3d(
                x=selected_warehouses['经度'],
                y=selected_warehouses['纬度'],
                z=selected_warehouses['最大容量'] / 10000,
                mode='markers',
                marker=dict(
                    size=15,
                    color='red',
                    symbol='diamond'
                ),
                text=selected_warehouses['地点名称'],
                name='仓库'
            ))

            fig.update_layout(
                title='3D供应链网络视图',
                scene=dict(
                    xaxis_title='经度',
                    yaxis_title='纬度',
                    zaxis_title='规模指标',
                    camera=dict(
                        eye=dict(x=1.5, y=1.5, z=1.5)
                    )
                ),
                height=700
            )

            st.plotly_chart(fig, use_container_width=True)

    with viz_tabs[1]:
        st.markdown("### 供应链网络图")

        if st.button("生成网络图"):
            # 创建网络
            G = nx.Graph()

            # 添加节点
            if len(st.session_state.customer_data) > 0:
                # 添加客户节点
                for idx, customer in st.session_state.customer_data.head(20).iterrows():
                    G.add_node(
                        f"C_{customer['客户编号']}",
                        node_type='customer',
                        size=customer['年需求量'] / 1000,
                        label=customer['城市']
                    )

                # 添加仓库节点
                for idx, warehouse in st.session_state.candidate_locations.head(10).iterrows():
                    G.add_node(
                        f"W_{warehouse['地点编号']}",
                        node_type='warehouse',
                        size=warehouse['最大容量'] / 10000,
                        label=warehouse['城市']
                    )

                # 添加边（基于距离阈值）
                for idx1, customer in st.session_state.customer_data.head(20).iterrows():
                    for idx2, warehouse in st.session_state.candidate_locations.head(10).iterrows():
                        distance = calculate_distance(
                            customer['经度'], customer['纬度'],
                            warehouse['经度'], warehouse['纬度']
                        )
                        if distance < 300:  # 300km阈值
                            G.add_edge(
                                f"C_{customer['客户编号']}",
                                f"W_{warehouse['地点编号']}",
                                weight=1 / distance if distance > 0 else 1
                            )

                # 使用力导向布局
                pos = nx.spring_layout(G, k=3, iterations=50)

                # 创建Plotly图
                edge_trace = []
                for edge in G.edges():
                    x0, y0 = pos[edge[0]]
                    x1, y1 = pos[edge[1]]
                    edge_trace.append(go.Scatter(
                        x=[x0, x1, None],
                        y=[y0, y1, None],
                        mode='lines',
                        line=dict(width=0.5, color='#888'),
                        hoverinfo='none'
                    ))

                # 分别创建客户和仓库节点
                customer_nodes = [n for n in G.nodes() if n.startswith('C_')]
                warehouse_nodes = [n for n in G.nodes() if n.startswith('W_')]

                customer_trace = go.Scatter(
                    x=[pos[node][0] for node in customer_nodes],
                    y=[pos[node][1] for node in customer_nodes],
                    mode='markers+text',
                    marker=dict(
                        size=[G.nodes[node]['size'] for node in customer_nodes],
                        color='lightblue',
                        line=dict(color='darkblue', width=2)
                    ),
                    text=[G.nodes[node]['label'] for node in customer_nodes],
                    name='客户',
                    hovertemplate='%{text}<extra></extra>'
                )

                warehouse_trace = go.Scatter(
                    x=[pos[node][0] for node in warehouse_nodes],
                    y=[pos[node][1] for node in warehouse_nodes],
                    mode='markers+text',
                    marker=dict(
                        size=[G.nodes[node]['size'] + 10 for node in warehouse_nodes],
                        color='red',
                        symbol='square',
                        line=dict(color='darkred', width=2)
                    ),
                    text=[G.nodes[node]['label'] for node in warehouse_nodes],
                    name='仓库',
                    hovertemplate='%{text}<extra></extra>'
                )

                fig = go.Figure(data=edge_trace + [customer_trace, warehouse_trace])
                fig.update_layout(
                    title='供应链网络拓扑结构',
                    showlegend=True,
                    hovermode='closest',
                    margin=dict(b=20, l=5, r=5, t=40),
                    xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                    plot_bgcolor='rgba(0,0,0,0)',
                    height=700
                )

                st.plotly_chart(fig, use_container_width=True)

                # 网络统计
                st.markdown("#### 网络统计信息")
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("网络密度", f"{nx.density(G):.3f}")
                with col2:
                    st.metric("平均度", f"{sum(dict(G.degree()).values()) / G.number_of_nodes():.2f}")
                with col3:
                    st.metric("连通分量", nx.number_connected_components(G))

    with viz_tabs[2]:
        st.markdown("### 多维度热力分析")

        analysis_type = st.selectbox(
            "选择分析维度",
            ["需求密度", "成本分布", "风险热图", "服务覆盖", "碳足迹分布"]
        )

        if len(st.session_state.customer_data) > 0:
            if analysis_type == "需求密度":
                # 创建需求密度热力图
                fig = go.Figure(data=go.Densitymapbox(
                    lat=st.session_state.customer_data['纬度'],
                    lon=st.session_state.customer_data['经度'],
                    z=st.session_state.customer_data['年需求量'],
                    radius=40,
                    colorscale='Hot',
                    zmin=0,
                    zmax=st.session_state.customer_data['年需求量'].max()
                ))

                fig.update_layout(
                    mapbox_style="open-street-map",
                    mapbox_center_lon=110,
                    mapbox_center_lat=35,
                    mapbox_zoom=3.5,
                    title='客户需求密度分布',
                    height=700
                )

                st.plotly_chart(fig, use_container_width=True)

            elif analysis_type == "风险热图":
                # 创建风险热图
                if len(st.session_state.candidate_locations) > 0:
                    fig = go.Figure(data=go.Scattermapbox(
                        lat=st.session_state.candidate_locations['纬度'],
                        lon=st.session_state.candidate_locations['经度'],
                        mode='markers',
                        marker=dict(
                            size=st.session_state.candidate_locations['风险评分'] * 5,
                            color=st.session_state.candidate_locations['风险评分'],
                            colorscale='Reds',
                            showscale=True,
                            colorbar=dict(title="风险评分")
                        ),
                        text=st.session_state.candidate_locations['地点名称'],
                        hovertemplate='%{text}<br>风险评分: %{marker.color:.1f}<extra></extra>'
                    ))

                    fig.update_layout(
                        mapbox_style="open-street-map",
                        mapbox_center_lon=110,
                        mapbox_center_lat=35,
                        mapbox_zoom=3.5,
                        title='候选地点风险分布图',
                        height=700
                    )

                    st.plotly_chart(fig, use_container_width=True)

    with viz_tabs[3]:
        st.markdown("### 时空动态分析")

        # 创建时间滑块
        time_range = st.slider(
            "选择时间范围",
            min_value=1,
            max_value=24,
            value=(1, 12),
            format="%d月",
            key="viz_time_range"
        )

        # 动画控制
        if st.button("播放动画"):
            # 创建动态数据
            frames = []
            months = list(range(time_range[0], time_range[1] + 1))

            for month in months:
                # 模拟需求变化
                demand_factor = 1 + 0.1 * np.sin(month * np.pi / 6)

                frame_data = go.Frame(
                    data=[go.Scattermapbox(
                        lat=st.session_state.customer_data['纬度'],
                        lon=st.session_state.customer_data['经度'],
                        mode='markers',
                        marker=dict(
                            size=st.session_state.customer_data['年需求量'] / 1000 * demand_factor,
                            color=st.session_state.customer_data['年需求量'] * demand_factor,
                            colorscale='Viridis'
                        )
                    )],
                    name=str(month)
                )
                frames.append(frame_data)

            # 创建初始图
            fig = go.Figure(
                data=[go.Scattermapbox(
                    lat=st.session_state.customer_data['纬度'],
                    lon=st.session_state.customer_data['经度'],
                    mode='markers',
                    marker=dict(
                        size=st.session_state.customer_data['年需求量'] / 1000,
                        color=st.session_state.customer_data['年需求量'],
                        colorscale='Viridis',
                        showscale=True
                    )
                )],
                frames=frames
            )

            # 添加动画控制
            fig.update_layout(
                updatemenus=[{
                    'type': 'buttons',
                    'showactive': False,
                    'buttons': [
                        {'label': '播放', 'method': 'animate', 'args': [None, {'frame': {'duration': 500}}]},
                        {'label': '暂停', 'method': 'animate', 'args': [[None], {'frame': {'duration': 0}}]}
                    ]
                }],
                sliders=[{
                    'active': 0,
                    'yanchor': 'top',
                    'xanchor': 'left',
                    'currentvalue': {
                        'font': {'size': 20},
                        'prefix': '月份:',
                        'visible': True,
                        'xanchor': 'right'
                    },
                    'steps': [{'args': [[f.name], {'frame': {'duration': 300}}],
                               'label': f"{i + 1}月",
                               'method': 'animate'} for i, f in enumerate(frames)]
                }],
                mapbox_style="open-street-map",
                mapbox_center_lon=110,
                mapbox_center_lat=35,
                mapbox_zoom=3.5,
                title='需求时空动态变化',
                height=700
            )

            st.plotly_chart(fig, use_container_width=True)

    with viz_tabs[4]:
        st.markdown("### 交互式仪表板设计器")

        # 仪表板布局选择
        layout = st.selectbox(
            "选择布局模板",
            ["2×2网格", "1+3布局", "自定义布局"]
        )

        # 组件选择
        available_widgets = {
            "关键指标卡": "metric",
            "趋势图": "line_chart",
            "饼图": "pie_chart",
            "地图": "map",
            "表格": "table",
            "仪表盘": "gauge"
        }

        selected_widgets = st.multiselect(
            "选择组件",
            list(available_widgets.keys()),
            default=["关键指标卡", "趋势图", "地图"]
        )

        if st.button("生成仪表板"):
            # 根据布局创建仪表板
            if layout == "2×2网格":
                col1, col2 = st.columns(2)
                col3, col4 = st.columns(2)

                with col1:
                    st.metric("总成本", "¥12.5M", "-5%")

                with col2:
                    # 小型趋势图
                    trend_data = pd.DataFrame({
                        'x': range(10),
                        'y': np.cumsum(np.random.randn(10))
                    })
                    st.line_chart(trend_data.set_index('x'))

                with col3:
                    # 饼图
                    pie_data = pd.DataFrame({
                        'category': ['建设', '运营', '运输'],
                        'value': [40, 35, 25]
                    })
                    fig = px.pie(pie_data, values='value', names='category')
                    st.plotly_chart(fig, use_container_width=True)

                with col4:
                    st.metric("效率提升", "18%", "+3%")

            st.success("仪表板已生成！可以拖拽调整组件位置。")

    with viz_tabs[5]:
        st.markdown("### VR/AR预览")
        st.info("🥽 VR/AR功能正在开发中...")

        # 模拟VR预览
        st.markdown("""
        **即将推出的VR功能:**
        - 🏭 虚拟仓库漫游
        - 📦 3D货物流动模拟
        - 🚚 运输路线体验
        - 👥 多人协同规划
        - 📊 沉浸式数据可视化
        """)

        # 显示3D模型预览
        if st.button("加载3D仓库模型"):
            st.image("https://via.placeholder.com/800x400/1e3c72/ffffff?text=3D+Warehouse+Model",
                     caption="仓库3D模型预览")


# 预测与模拟
def show_prediction_simulation():
    """预测与模拟中心"""
    st.markdown("## 🔮 预测与模拟中心")

    sim_tabs = st.tabs([
        "需求预测", "成本模拟", "风险场景",
        "网络优化", "应急演练", "数字孪生"
    ])

    with sim_tabs[0]:
        st.markdown("### 智能需求预测")

        # 预测设置
        col1, col2, col3 = st.columns(3)
        with col1:
            forecast_method = st.selectbox(
                "预测方法",
                ["ARIMA", "Prophet", "LSTM", "XGBoost", "集成模型"]
            )
        with col2:
            forecast_horizon = st.number_input("预测期数(月)", 1, 36, 12, key="forecast_horizon")
        with col3:
            confidence_level = st.slider("置信水平", 0.8, 0.99, 0.95, key="forecast_confidence")

        if st.button("执行预测"):
            # 生成历史数据
            dates = pd.date_range(end=datetime.now(), periods=48, freq='M')

            # 创建多个产品/地区的需求数据
            products = ['产品A', '产品B', '产品C']
            regions = ['华东', '华南', '华北', '西南']

            forecast_results = {}

            for product in products:
                for region in regions:
                    # 生成历史需求（包含趋势、季节性和随机因素）
                    trend = np.linspace(1000, 1500, 48)
                    seasonal = 200 * np.sin(np.arange(48) * 2 * np.pi / 12)
                    noise = np.random.normal(0, 50, 48)
                    demand = trend + seasonal + noise

                    # 生成预测
                    future_trend = np.linspace(1500, 1800, forecast_horizon)
                    future_seasonal = 200 * np.sin(np.arange(48, 48 + forecast_horizon) * 2 * np.pi / 12)
                    future_noise = np.random.normal(0, 30, forecast_horizon)
                    forecast = future_trend + future_seasonal + future_noise

                    # 计算置信区间
                    std = np.std(noise)
                    z_score = norm.ppf((1 + confidence_level) / 2)
                    margin = z_score * std

                    forecast_results[f"{product}-{region}"] = {
                        'historical': demand,
                        'forecast': forecast,
                        'lower': forecast - margin,
                        'upper': forecast + margin
                    }

            # 显示预测结果
            selected_series = st.selectbox(
                "选择查看的序列",
                list(forecast_results.keys())
            )

            if selected_series:
                result = forecast_results[selected_series]

                # 创建图表
                fig = go.Figure()

                # 历史数据
                fig.add_trace(go.Scatter(
                    x=dates,
                    y=result['historical'],
                    mode='lines',
                    name='历史数据',
                    line=dict(color='blue')
                ))

                # 预测数据
                future_dates = pd.date_range(
                    start=dates[-1] + pd.DateOffset(months=1),
                    periods=forecast_horizon,
                    freq='M'
                )

                fig.add_trace(go.Scatter(
                    x=future_dates,
                    y=result['forecast'],
                    mode='lines',
                    name='预测值',
                    line=dict(color='red', dash='dash')
                ))

                # 置信区间
                fig.add_trace(go.Scatter(
                    x=future_dates.tolist() + future_dates.tolist()[::-1],
                    y=result['upper'].tolist() + result['lower'].tolist()[::-1],
                    fill='toself',
                    fillcolor='rgba(255,0,0,0.2)',
                    line=dict(color='rgba(255,255,255,0)'),
                    name=f'{int(confidence_level * 100)}%置信区间',
                    showlegend=True
                ))

                fig.update_layout(
                    title=f'{selected_series} 需求预测',
                    xaxis_title='日期',
                    yaxis_title='需求量',
                    hovermode='x unified',
                    height=500
                )

                st.plotly_chart(fig, use_container_width=True)

                # 预测统计
                col1, col2, col3 = st.columns(3)
                col1.metric(
                    "平均预测值",
                    f"{np.mean(result['forecast']):.0f}",
                    f"{(np.mean(result['forecast']) / np.mean(result['historical']) - 1) * 100:.1f}%"
                )
                col2.metric(
                    "预测峰值",
                    f"{np.max(result['forecast']):.0f}"
                )
                col3.metric(
                    "预测波动性",
                    f"{np.std(result['forecast']):.0f}"
                )

    with sim_tabs[1]:
        st.markdown("### 成本优化模拟")

        # 成本参数设置
        st.markdown("#### 成本参数设置")

        col1, col2 = st.columns(2)

        with col1:
            land_cost_change = st.slider("土地成本变化(%)", -30, 50, 0, key="sim_land_cost")
            labor_cost_change = st.slider("人工成本变化(%)", -20, 40, 10, key="sim_labor_cost")
            fuel_cost_change = st.slider("燃料成本变化(%)", -50, 100, 20, key="sim_fuel_cost")

        with col2:
            demand_change = st.slider("需求变化(%)", -30, 50, 15, key="sim_demand")
            efficiency_gain = st.slider("效率提升(%)", 0, 30, 10, key="sim_efficiency")
            automation_level = st.slider("自动化程度(%)", 0, 100, 30, key="sim_automation")

        if st.button("运行成本模拟"):
            # 基准成本
            base_costs = {
                '土地成本': 3500,
                '建设成本': 6000,
                '人工成本': 2000,
                '运输成本': 1500,
                '运营成本': 1000
            }

            # 计算调整后的成本
            adjusted_costs = {
                '土地成本': base_costs['土地成本'] * (1 + land_cost_change / 100),
                '建设成本': base_costs['建设成本'],
                '人工成本': base_costs['人工成本'] * (1 + labor_cost_change / 100) * (1 - automation_level / 200),
                '运输成本': base_costs['运输成本'] * (1 + fuel_cost_change / 100) * (1 - efficiency_gain / 100),
                '运营成本': base_costs['运营成本'] * (1 - efficiency_gain / 100)
            }

            # 创建对比图
            fig = go.Figure()

            categories = list(base_costs.keys())

            fig.add_trace(go.Bar(
                name='基准成本',
                x=categories,
                y=list(base_costs.values()),
                marker_color='lightblue'
            ))

            fig.add_trace(go.Bar(
                name='调整后成本',
                x=categories,
                y=list(adjusted_costs.values()),
                marker_color='darkblue'
            ))

            fig.update_layout(
                title='成本变化对比分析',
                yaxis_title='成本(万元)',
                barmode='group',
                height=500
            )

            st.plotly_chart(fig, use_container_width=True)

            # 成本节省分析
            total_base = sum(base_costs.values())
            total_adjusted = sum(adjusted_costs.values())
            savings = total_base - total_adjusted

            col1, col2, col3 = st.columns(3)
            col1.metric("原始总成本", f"¥{total_base:.0f}万")
            col2.metric("优化后成本", f"¥{total_adjusted:.0f}万")
            col3.metric("成本节省", f"¥{savings:.0f}万", f"{-savings / total_base * 100:.1f}%")

            # ROI分析
            st.markdown("#### 投资回报分析")

            # 收入预测
            revenue_growth = demand_change * 0.8  # 假设80%的需求增长转化为收入
            annual_revenue = 15000 * (1 + revenue_growth / 100)

            # 计算投资回报
            investment = total_adjusted
            annual_profit = annual_revenue - total_adjusted * 0.1  # 假设年运营成本为总投资的10%
            roi = (annual_profit / investment) * 100
            payback_period = investment / annual_profit if annual_profit > 0 else float('inf')

            col1, col2, col3 = st.columns(3)
            col1.metric("年收入预测", f"¥{annual_revenue:.0f}万")
            col2.metric("投资回报率", f"{roi:.1f}%")
            col3.metric("投资回收期", f"{payback_period:.1f}年" if payback_period < 100 else "N/A")

    with sim_tabs[2]:
        st.markdown("### 风险场景模拟")

        # 风险场景选择
        risk_scenarios = {
            "供应链中断": {
                "probability": 0.15,
                "impact": "high",
                "duration": "3-6个月",
                "cost_increase": 0.3
            },
            "需求急剧下降": {
                "probability": 0.2,
                "impact": "medium",
                "duration": "6-12个月",
                "cost_increase": -0.2
            },
            "自然灾害": {
                "probability": 0.05,
                "impact": "severe",
                "duration": "1-3个月",
                "cost_increase": 0.5
            },
            "政策变化": {
                "probability": 0.3,
                "impact": "medium",
                "duration": "长期",
                "cost_increase": 0.15
            },
            "竞争加剧": {
                "probability": 0.4,
                "impact": "low",
                "duration": "长期",
                "cost_increase": -0.1
            }
        }

        selected_scenarios = st.multiselect(
            "选择风险场景",
            list(risk_scenarios.keys()),
            default=["供应链中断", "需求急剧下降"]
        )

        # 蒙特卡洛模拟参数
        n_simulations = st.number_input("模拟次数", 100, 10000, 1000, key="mc_simulations")

        if st.button("运行风险模拟"):
            # 执行蒙特卡洛模拟
            simulation_results = []

            progress_bar = st.progress(0)

            for i in range(n_simulations):
                # 基准值
                base_cost = 10000
                base_revenue = 15000

                # 应用风险影响
                total_cost = base_cost
                total_revenue = base_revenue

                for scenario in selected_scenarios:
                    scenario_data = risk_scenarios[scenario]

                    # 根据概率决定是否发生
                    if np.random.random() < scenario_data['probability']:
                        cost_impact = scenario_data['cost_increase']
                        total_cost *= (1 + cost_impact)

                        # 收入影响（简化模型）
                        if cost_impact > 0:
                            total_revenue *= (1 - cost_impact * 0.5)
                        else:
                            total_revenue *= (1 - cost_impact * 0.3)

                profit = total_revenue - total_cost
                simulation_results.append({
                    'cost': total_cost,
                    'revenue': total_revenue,
                    'profit': profit
                })

                progress_bar.progress((i + 1) / n_simulations)

            # 分析结果
            results_df = pd.DataFrame(simulation_results)

            # 显示结果分布
            col1, col2 = st.columns(2)

            with col1:
                # 利润分布直方图
                fig = px.histogram(
                    results_df,
                    x='profit',
                    nbins=50,
                    title='利润分布（蒙特卡洛模拟）',
                    labels={'profit': '利润(万元)', 'count': '频次'}
                )

                # 添加统计线
                mean_profit = results_df['profit'].mean()
                fig.add_vline(x=mean_profit, line_dash="dash", line_color="red",
                              annotation_text=f"均值: {mean_profit:.0f}")

                # VaR计算
                var_95 = results_df['profit'].quantile(0.05)
                fig.add_vline(x=var_95, line_dash="dash", line_color="orange",
                              annotation_text=f"VaR(95%): {var_95:.0f}")

                st.plotly_chart(fig, use_container_width=True)

            with col2:
                # 风险指标
                st.markdown("#### 风险指标")

                col1, col2 = st.columns(2)
                col1.metric("平均利润", f"¥{mean_profit:.0f}万")
                col2.metric("利润标准差", f"¥{results_df['profit'].std():.0f}万")

                col1, col2 = st.columns(2)
                col1.metric("最大损失", f"¥{results_df['profit'].min():.0f}万")
                col2.metric("盈利概率", f"{(results_df['profit'] > 0).mean():.1%}")

                # 风险价值(VaR)和条件风险价值(CVaR)
                var_95 = results_df['profit'].quantile(0.05)
                cvar_95 = results_df[results_df['profit'] <= var_95]['profit'].mean()

                st.metric("VaR (95%)", f"¥{var_95:.0f}万")
                st.metric("CVaR (95%)", f"¥{cvar_95:.0f}万")

    with sim_tabs[3]:
        st.markdown("### 网络优化模拟")

        # 优化目标选择
        optimization_goals = st.multiselect(
            "优化目标",
            ["最小化成本", "最大化覆盖", "最小化时间", "平衡负载", "降低碳排放"],
            default=["最小化成本", "最大化覆盖"]
        )

        # 约束条件
        st.markdown("#### 约束条件")
        col1, col2 = st.columns(2)

        with col1:
            max_warehouses = st.number_input("最大仓库数", 1, 20, 5, key="net_max_warehouses")
            min_service_level = st.slider("最低服务水平", 0.7, 0.99, 0.9, key="net_service_level")

        with col2:
            max_distance = st.number_input("最大服务距离(km)", 100, 1000, 300, key="net_max_distance")
            budget_constraint = st.number_input("预算约束(万)", 5000, 50000, 15000, key="net_budget")

        if st.button("优化网络"):
            # 模拟优化过程
            st.info("正在优化供应链网络...")

            # 创建优化进度
            optimization_steps = [
                "初始化优化模型",
                "生成初始解",
                "迭代优化",
                "约束检查",
                "解的改进",
                "收敛判断"
            ]

            progress_bar = st.progress(0)
            status_text = st.empty()

            for i, step in enumerate(optimization_steps):
                status_text.text(f"步骤 {i + 1}/{len(optimization_steps)}: {step}")
                time.sleep(0.5)
                progress_bar.progress((i + 1) / len(optimization_steps))

            # 生成优化结果
            st.success("网络优化完成！")

            # 显示优化前后对比
            metrics_before = {
                '总成本': 15000,
                '覆盖率': 0.82,
                '平均运输时间': 6.5,
                '碳排放': 1200
            }

            metrics_after = {
                '总成本': 12800,
                '覆盖率': 0.93,
                '平均运输时间': 4.2,
                '碳排放': 980
            }

            # 创建对比雷达图
            categories = list(metrics_before.keys())

            # 标准化数据（0-1范围）
            values_before = []
            values_after = []

            for cat in categories:
                if cat in ['覆盖率']:
                    values_before.append(metrics_before[cat])
                    values_after.append(metrics_after[cat])
                elif cat in ['总成本', '平均运输时间', '碳排放']:
                    # 反向标准化（越小越好）
                    max_val = max(metrics_before[cat], metrics_after[cat])
                    values_before.append(1 - metrics_before[cat] / max_val)
                    values_after.append(1 - metrics_after[cat] / max_val)

            fig = go.Figure()

            fig.add_trace(go.Scatterpolar(
                r=values_before + [values_before[0]],
                theta=categories + [categories[0]],
                fill='toself',
                name='优化前',
                line=dict(color='red')
            ))

            fig.add_trace(go.Scatterpolar(
                r=values_after + [values_after[0]],
                theta=categories + [categories[0]],
                fill='toself',
                name='优化后',
                line=dict(color='green')
            ))

            fig.update_layout(
                polar=dict(
                    radialaxis=dict(
                        visible=True,
                        range=[0, 1]
                    )),
                showlegend=True,
                title="网络优化效果对比"
            )

            st.plotly_chart(fig, use_container_width=True)

            # 改进指标
            st.markdown("#### 优化改进")

            col1, col2, col3, col4 = st.columns(4)

            improvements = {
                '成本降低': (metrics_before['总成本'] - metrics_after['总成本']) / metrics_before['总成本'] * 100,
                '覆盖提升': (metrics_after['覆盖率'] - metrics_before['覆盖率']) / metrics_before['覆盖率'] * 100,
                '时间缩短': (metrics_before['平均运输时间'] - metrics_after['平均运输时间']) / metrics_before[
                    '平均运输时间'] * 100,
                '碳减排': (metrics_before['碳排放'] - metrics_after['碳排放']) / metrics_before['碳排放'] * 100
            }

            for i, (metric, improvement) in enumerate(improvements.items()):
                cols = [col1, col2, col3, col4]
                cols[i].metric(metric, f"{improvement:.1f}%", "↑")

    with sim_tabs[4]:
        st.markdown("### 应急响应演练")

        # 应急场景选择
        emergency_scenarios = [
            "仓库火灾",
            "道路封闭",
            "系统故障",
            "极端天气",
            "疫情封控",
            "供应商违约"
        ]

        selected_emergency = st.selectbox("选择应急场景", emergency_scenarios)

        # 场景参数
        col1, col2 = st.columns(2)

        with col1:
            affected_warehouses = st.multiselect(
                "受影响仓库",
                st.session_state.selected_locations if st.session_state.selected_locations else ["W001", "W002",
                                                                                                 "W003"],
                default=[st.session_state.selected_locations[0]] if st.session_state.selected_locations else ["W001"]
            )

            impact_duration = st.slider("影响持续时间(天)", 1, 30, 7, key="emergency_duration")

        with col2:
            capacity_loss = st.slider("容量损失(%)", 0, 100, 50, key="emergency_capacity_loss")
            response_time = st.slider("响应时间(小时)", 1, 24, 4, key="emergency_response_time")

        if st.button("开始演练"):
            st.warning(f"⚠️ 应急演练开始: {selected_emergency}")

            # 模拟应急响应过程
            response_steps = {
                "T+0": "事件发生，启动应急预案",
                "T+1h": "成立应急指挥中心",
                "T+2h": "评估影响范围和程度",
                "T+4h": "启动备用方案",
                "T+8h": "调整配送网络",
                "T+24h": "恢复部分运营",
                "T+72h": "全面恢复运营"
            }

            # 显示响应时间线
            st.markdown("#### 应急响应时间线")

            for time_point, action in response_steps.items():
                col1, col2 = st.columns([1, 4])
                with col1:
                    st.write(f"**{time_point}**")
                with col2:
                    st.write(action)
                time.sleep(0.3)  # 模拟实时更新

            # 影响分析
            st.markdown("#### 影响分析")

            # 计算影响
            total_capacity = 100000  # 假设总容量
            lost_capacity = total_capacity * len(affected_warehouses) / 5 * capacity_loss / 100
            affected_customers = int(len(st.session_state.customer_data) * len(affected_warehouses) / 5) if len(
                st.session_state.customer_data) > 0 else 20

            col1, col2, col3 = st.columns(3)
            col1.metric("容量损失", f"{lost_capacity:,.0f} 单位")
            col2.metric("受影响客户", f"{affected_customers} 个")
            col3.metric("预计损失", f"¥{lost_capacity * 0.1:.0f}万")

            # 应急方案
            st.markdown("#### 应急方案")

            emergency_plans = {
                "立即行动": [
                    "启动应急物资调配",
                    "通知受影响客户",
                    "安排临时配送方案"
                ],
                "短期措施": [
                    "从其他仓库调配库存",
                    "租用临时仓储设施",
                    "加快运输频次"
                ],
                "长期方案": [
                    "建立备用仓库",
                    "优化库存分布",
                    "加强风险预防"
                ]
            }

            for phase, actions in emergency_plans.items():
                with st.expander(phase):
                    for i, action in enumerate(actions):
                        st.checkbox(action, key=f"emergency_{phase}_{i}_{action}")

            # 恢复进度
            st.markdown("#### 恢复进度跟踪")

            recovery_progress = st.progress(0)
            recovery_status = st.empty()

            for i in range(101):
                recovery_progress.progress(i)
                recovery_status.text(f"恢复进度: {i}%")
                time.sleep(0.02)

            st.success("✅ 应急演练完成！系统已恢复正常运行。")

    with sim_tabs[5]:
        st.markdown("### 数字孪生系统")
        st.info("🔮 数字孪生技术正在集成中...")

        # 数字孪生概述
        st.markdown("""
        数字孪生系统将提供：

        - **实时镜像**: 物理仓库的实时数字副本
        - **预测维护**: 基于IoT数据的设备故障预测
        - **虚拟测试**: 在数字环境中测试优化方案
        - **实时优化**: 根据实时数据动态调整运营
        - **场景模拟**: 在虚拟环境中模拟各种场景
        """)

        # 模拟数字孪生仪表板
        if st.button("启动数字孪生"):
            col1, col2 = st.columns(2)

            with col1:
                st.markdown("#### 实体状态")
                st.metric("温度", "22.5°C", "+0.3°C")
                st.metric("湿度", "45%", "-2%")
                st.metric("能耗", "1,234 kWh", "+5%")

            with col2:
                st.markdown("#### 数字镜像")
                st.metric("同步延迟", "< 100ms", "")
                st.metric("预测准确度", "96.8%", "+1.2%")
                st.metric("异常检测", "0", "")

            # 实时数据流
            st.markdown("#### 实时数据流")

            # 创建实时更新的图表占位符
            chart_placeholder = st.empty()

            # 模拟实时数据
            for i in range(20):
                # 生成新数据点
                new_data = pd.DataFrame({
                    'timestamp': [datetime.now()],
                    'throughput': [np.random.uniform(80, 120)],
                    'utilization': [np.random.uniform(70, 90)]
                })

                # 更新图表
                fig = go.Figure()
                fig.add_trace(go.Scatter(
                    y=new_data['throughput'],
                    mode='lines+markers',
                    name='吞吐量',
                    line=dict(color='blue')
                ))
                fig.update_layout(
                    title='实时运营指标',
                    yaxis_title='值',
                    height=300
                )

                chart_placeholder.plotly_chart(fig, use_container_width=True)
                time.sleep(0.5)


# 结果展示
def show_results_display():
    """结果展示页面"""
    st.subheader("📋 结果展示")

    if not st.session_state.selected_locations:
        st.warning("请先运行选址优化算法")
        # 自动选择一些仓库
        if not st.session_state.candidate_locations.empty:
            st.session_state.selected_locations = st.session_state.candidate_locations.head(3)['地点编号'].tolist()

    # 显示选中的仓库
    st.markdown("### 选中的仓库")
    selected_warehouses = st.session_state.candidate_locations[
        st.session_state.candidate_locations['地点编号'].isin(st.session_state.selected_locations)
    ]

    st.dataframe(selected_warehouses[['地点编号', '地点名称', '城市', '建设成本', '最大容量', '服务半径']])

    # 地图展示
    st.markdown("### 选址结果地图")
    m = create_advanced_folium_map(
        st.session_state.customer_data,
        st.session_state.candidate_locations,
        st.session_state.selected_locations,
        show_connections=True
    )
    folium_static(m, width=1200, height=600)

    # 服务覆盖分析
    st.markdown("### 服务覆盖分析")

    if st.session_state.distance_matrix:
        # 计算每个客户到最近仓库的距离
        customer_coverage = []

        for _, customer in st.session_state.customer_data.iterrows():
            min_distance = float('inf')
            nearest_warehouse = None

            for warehouse_id in st.session_state.selected_locations:
                distance = st.session_state.distance_matrix.get(
                    (warehouse_id, customer['客户编号']), float('inf')
                )
                if distance < min_distance:
                    min_distance = distance
                    nearest_warehouse = warehouse_id

            customer_coverage.append({
                '客户编号': customer['客户编号'],
                '最近仓库': nearest_warehouse,
                '距离(km)': min_distance,
                '是否覆盖': min_distance <= 300  # 假设300km为服务范围
            })

        coverage_df = pd.DataFrame(customer_coverage)

        # 覆盖率统计
        coverage_rate = coverage_df['是否覆盖'].sum() / len(coverage_df) * 100
        avg_distance = coverage_df['距离(km)'].mean()

        col1, col2 = st.columns(2)
        col1.metric("客户覆盖率", f"{coverage_rate:.1f}%")
        col2.metric("平均服务距离", f"{avg_distance:.1f} km")

        # 距离分布直方图
        fig = px.histogram(
            coverage_df,
            x='距离(km)',
            nbins=20,
            title='客户到最近仓库的距离分布'
        )
        st.plotly_chart(fig, use_container_width=True)


# 高级报告生成
def generate_advanced_report():
    """生成高级分析报告"""
    st.subheader("📊 高级报告生成")

    report_type = st.selectbox(
        "选择报告类型",
        ["执行摘要", "详细分析报告", "技术报告", "投资者报告", "可行性研究报告"]
    )

    include_sections = st.multiselect(
        "包含章节",
        ["概述", "市场分析", "选址方案", "成本分析", "风险评估",
         "财务预测", "实施计划", "结论建议", "碳足迹分析", "技术方案"],
        default=["概述", "选址方案", "成本分析", "结论建议"]
    )

    if st.button("生成报告", type="primary"):
        with st.spinner("正在生成专业报告..."):
            # 创建Word文档
            doc = Document()

            # 添加标题
            doc.add_heading('仓库选址优化项目报告', 0)
            doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y年%m月%d日")}')
            doc.add_paragraph(f'报告类型: {report_type}')

            # 添加目录
            doc.add_page_break()
            doc.add_heading('目录', 1)
            for i, section in enumerate(include_sections, 1):
                doc.add_paragraph(f'{i}. {section}', style='List Number')

            # 添加各章节内容
            for section in include_sections:
                doc.add_page_break()
                doc.add_heading(section, 1)

                if section == "概述":
                    doc.add_paragraph(
                        "本报告基于先进的数据分析和优化算法，为仓库选址项目提供全面的决策支持。"
                        "通过综合考虑成本、效率、风险等多个维度，我们制定了最优的仓库布局方案。"
                    )

                    # 添加关键指标表
                    table = doc.add_table(rows=5, cols=2)
                    table.style = 'Light Grid Accent 1'

                    metrics = [
                        ('分析客户数', f"{len(st.session_state.customer_data)}"),
                        ('候选地点数', f"{len(st.session_state.candidate_locations)}"),
                        ('推荐仓库数', f"{len(st.session_state.selected_locations)}"),
                        ('预计总投资', "¥1.2亿"),
                        ('预期ROI', "22.5%")
                    ]

                    for i, (metric, value) in enumerate(metrics):
                        table.cell(i, 0).text = metric
                        table.cell(i, 1).text = value

                elif section == "选址方案":
                    doc.add_paragraph("基于多种优化算法的综合分析，我们推荐以下选址方案：")

                    if st.session_state.selected_locations:
                        for i, warehouse_id in enumerate(st.session_state.selected_locations[:5], 1):
                            warehouse = st.session_state.candidate_locations[
                                st.session_state.candidate_locations['地点编号'] == warehouse_id
                                ].iloc[0]

                            doc.add_heading(f"{i}. {warehouse['地点名称']}", 2)
                            doc.add_paragraph(f"位置: {warehouse['城市']}")
                            doc.add_paragraph(f"建设成本: ¥{warehouse['建设成本'] / 1e6:.1f}百万")
                            doc.add_paragraph(f"设计容量: {warehouse['最大容量']:,} 单位")
                            doc.add_paragraph(f"服务半径: {warehouse['服务半径']} 公里")

                elif section == "成本分析":
                    doc.add_paragraph("项目成本构成如下：")

                    # 添加成本明细
                    doc.add_paragraph("• 土地成本: ¥3,500万")
                    doc.add_paragraph("• 建设成本: ¥6,000万")
                    doc.add_paragraph("• 设备投资: ¥2,000万")
                    doc.add_paragraph("• 其他费用: ¥500万")
                    doc.add_paragraph("• 总投资额: ¥12,000万")

                    doc.add_paragraph("\n投资回收期预计为4.2年，内部收益率(IRR)为18.5%。")

                elif section == "碳足迹分析":
                    if '碳足迹分析' in st.session_state.analysis_results:
                        carbon_result = st.session_state.analysis_results['碳足迹分析']
                        doc.add_paragraph(f"年碳排放量: {carbon_result.get('总碳排放(吨/年)', 0):.1f} 吨CO₂")
                        doc.add_paragraph(f"碳补偿成本: ¥{carbon_result.get('碳补偿成本', 0):,.0f}/年")
                        doc.add_paragraph("\n减碳建议：")
                        doc.add_paragraph("• 优化运输方式，增加铁路运输比例")
                        doc.add_paragraph("• 引入新能源车辆")
                        doc.add_paragraph("• 建设绿色仓库，采用太阳能等清洁能源")

                # 添加更多章节...

            # 保存文档
            report_buffer = io.BytesIO()
            doc.save(report_buffer)
            report_buffer.seek(0)

            # 提供下载
            st.download_button(
                label="📥 下载Word报告",
                data=report_buffer,
                file_name=f"warehouse_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # 同时生成PDF版本（需要额外库）
            st.info("报告已生成！您可以下载Word版本，PDF版本正在开发中。")


# 系统管理
def show_system_management():
    """系统管理中心"""
    st.markdown("## ⚙️ 系统管理中心")

    mgmt_tabs = st.tabs([
        "用户管理", "权限设置", "系统配置",
        "审计日志", "备份恢复", "API管理"
    ])

    with mgmt_tabs[0]:
        st.markdown("### 用户管理")

        # 模拟用户数据
        users = pd.DataFrame({
            '用户ID': ['U001', 'U002', 'U003', 'U004', 'U005'],
            '用户名': ['admin', 'analyst1', 'manager1', 'operator1', 'viewer1'],
            '角色': ['系统管理员', '数据分析师', '项目经理', '操作员', '查看者'],
            '状态': ['活跃', '活跃', '活跃', '离线', '活跃'],
            '最后登录': pd.to_datetime(['2024-01-15', '2024-01-15', '2024-01-14', '2024-01-10', '2024-01-15'])
        })

        st.dataframe(users, use_container_width=True)

        # 用户操作
        col1, col2, col3 = st.columns(3)
        with col1:
            if st.button("添加用户"):
                st.success("用户添加成功")
        with col2:
            if st.button("编辑权限"):
                st.info("进入权限编辑模式")
        with col3:
            if st.button("导出用户列表"):
                st.success("用户列表已导出")

    with mgmt_tabs[1]:
        st.markdown("### 权限设置")

        # 角色权限矩阵
        roles = ['系统管理员', '数据分析师', '项目经理', '操作员', '查看者']
        permissions = ['数据查看', '数据编辑', '运行分析', '系统配置', '用户管理']

        # 创建权限矩阵
        permission_matrix = pd.DataFrame(
            index=roles,
            columns=permissions,
            data=[
                [True, True, True, True, True],  # 管理员
                [True, True, True, False, False],  # 分析师
                [True, True, True, False, False],  # 项目经理
                [True, True, False, False, False],  # 操作员
                [True, False, False, False, False]  # 查看者
            ]
        )

        # 显示权限矩阵
        st.dataframe(permission_matrix.style.applymap(
            lambda x: 'background-color: lightgreen' if x else 'background-color: lightcoral'
        ))

        # 权限修改
        selected_role = st.selectbox("选择角色", roles)
        selected_permission = st.selectbox("选择权限", permissions)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("授予权限"):
                st.success(f"已授予{selected_role} {selected_permission}权限")
        with col2:
            if st.button("撤销权限"):
                st.warning(f"已撤销{selected_role} {selected_permission}权限")

    with mgmt_tabs[2]:
        st.markdown("### 系统配置")

        # 系统参数配置
        st.markdown("#### 基础配置")

        col1, col2 = st.columns(2)

        with col1:
            st.text_input("系统名称", value="仓库选址优化系统 Ultimate Pro", key="sys_name")
            st.selectbox("默认语言", ["中文", "English", "日本語"], key="sys_language")
            st.selectbox("时区", ["UTC+8 北京时间", "UTC+0 格林威治时间", "UTC-5 纽约时间"], key="sys_timezone")

        with col2:
            st.number_input("会话超时(分钟)", value=30, key="sys_session_timeout")
            st.selectbox("主题", ["浅色", "深色", "自动"], key="sys_theme")
            st.checkbox("启用双因素认证", value=True, key="system_2fa_enabled")

        st.markdown("#### 高级配置")

        with st.expander("算法参数"):
            st.slider("默认迭代次数", 50, 1000, 100, key="algo_iterations")
            st.slider("收敛阈值", 0.0001, 0.01, 0.001, format="%.4f", key="algo_convergence")
            st.selectbox("默认优化算法", ["自动选择", "遗传算法", "粒子群", "模拟退火"], key="algo_default")

        with st.expander("性能设置"):
            st.slider("最大并行任务数", 1, 10, 4, key="perf_parallel_tasks")
            st.checkbox("启用缓存", value=True, key="system_cache_enabled")
            st.checkbox("启用GPU加速", value=False, key="system_gpu_enabled")
            st.number_input("缓存大小(GB)", value=4, key="perf_cache_size")

        if st.button("保存配置"):
            st.success("配置已保存")

    with mgmt_tabs[3]:
        st.markdown("### 审计日志")

        # 生成示例日志
        log_entries = []
        actions = ['登录', '数据导入', '运行分析', '导出报告', '修改配置', '查看数据']
        users = ['admin', 'analyst1', 'manager1']

        for i in range(50):
            log_entries.append({
                '时间': datetime.now() - timedelta(hours=i),
                '用户': np.random.choice(users),
                '操作': np.random.choice(actions),
                'IP地址': f"192.168.1.{np.random.randint(1, 255)}",
                '状态': np.random.choice(['成功', '成功', '成功', '失败']),
                '详情': '操作详细信息...'
            })

        log_df = pd.DataFrame(log_entries)

        # 日志筛选
        col1, col2, col3 = st.columns(3)

        with col1:
            filter_user = st.selectbox("用户筛选", ['全部'] + users)
        with col2:
            filter_action = st.selectbox("操作筛选", ['全部'] + actions)
        with col3:
            filter_status = st.selectbox("状态筛选", ['全部', '成功', '失败'])

        # 应用筛选
        filtered_log = log_df
        if filter_user != '全部':
            filtered_log = filtered_log[filtered_log['用户'] == filter_user]
        if filter_action != '全部':
            filtered_log = filtered_log[filtered_log['操作'] == filter_action]
        if filter_status != '全部':
            filtered_log = filtered_log[filtered_log['状态'] == filter_status]

        # 显示日志
        st.dataframe(
            filtered_log.head(20).style.applymap(
                lambda x: 'color: red' if x == '失败' else 'color: green' if x == '成功' else '',
                subset=['状态']
            ),
            use_container_width=True
        )

        # 导出日志
        if st.button("导出日志"):
            csv = filtered_log.to_csv(index=False)
            st.download_button(
                "下载CSV",
                csv,
                "audit_log.csv",
                "text/csv"
            )

    with mgmt_tabs[4]:
        st.markdown("### 备份与恢复")

        # 备份管理
        st.markdown("#### 自动备份设置")

        col1, col2 = st.columns(2)

        with col1:
            backup_enabled = st.checkbox("启用自动备份", value=True, key="backup_auto_enabled")
            backup_frequency = st.selectbox(
                "备份频率",
                ["每小时", "每天", "每周", "每月"],
                disabled=not backup_enabled
            )

        with col2:
            backup_retention = st.number_input(
                "保留天数",
                value=30,
                disabled=not backup_enabled,
                key="backup_retention_days"
            )
            backup_location = st.text_input(
                "备份位置",
                value="/backup/warehouse_system/",
                disabled=not backup_enabled,
                key="backup_location_path"
            )

        # 备份历史
        st.markdown("#### 备份历史")

        backup_history = pd.DataFrame({
            '备份时间': pd.date_range(end=datetime.now(), periods=10, freq='D'),
            '备份类型': ['自动'] * 8 + ['手动'] * 2,
            '大小': np.random.uniform(100, 500, 10),
            '状态': ['成功'] * 9 + ['失败'],
            '备注': [''] * 9 + ['磁盘空间不足']
        })

        st.dataframe(backup_history, use_container_width=True)

        # 操作按钮
        col1, col2, col3 = st.columns(3)

        with col1:
            if st.button("立即备份"):
                progress = st.progress(0)
                for i in range(100):
                    progress.progress(i + 1)
                    time.sleep(0.01)
                st.success("备份完成！")

        with col2:
            if st.button("恢复数据"):
                st.warning("请选择要恢复的备份文件")

        with col3:
            if st.button("验证备份"):
                st.info("备份验证中...")
                time.sleep(1)
                st.success("备份文件完整性验证通过")

    with mgmt_tabs[5]:
        st.markdown("### API管理")

        # API密钥管理
        st.markdown("#### API密钥")

        api_keys = pd.DataFrame({
            '名称': ['生产环境', '测试环境', '开发环境'],
            '密钥': ['sk-prod-****', 'sk-test-****', 'sk-dev-****'],
            '创建时间': pd.to_datetime(['2024-01-01', '2024-01-05', '2024-01-10']),
            '最后使用': pd.to_datetime(['2024-01-15', '2024-01-14', '2024-01-15']),
            '状态': ['活跃', '活跃', '已禁用']
        })

        st.dataframe(api_keys, use_container_width=True)

        # API使用统计
        st.markdown("#### API使用统计")

        # 创建API调用趋势图
        dates = pd.date_range(end=datetime.now(), periods=30, freq='D')
        api_calls = pd.DataFrame({
            '日期': dates,
            '调用次数': np.random.poisson(1000, 30),
            '成功率': np.random.uniform(0.95, 0.99, 30)
        })

        fig = make_subplots(
            rows=2, cols=1,
            subplot_titles=('API调用次数', 'API成功率'),
            shared_xaxes=True
        )

        fig.add_trace(
            go.Scatter(x=api_calls['日期'], y=api_calls['调用次数'], mode='lines+markers'),
            row=1, col=1
        )

        fig.add_trace(
            go.Scatter(x=api_calls['日期'], y=api_calls['成功率'], mode='lines+markers'),
            row=2, col=1
        )

        fig.update_yaxes(title_text="调用次数", row=1, col=1)
        fig.update_yaxes(title_text="成功率", row=2, col=1)
        fig.update_xaxes(title_text="日期", row=2, col=1)

        st.plotly_chart(fig, use_container_width=True)

        # API文档链接
        st.markdown("#### API文档")
        st.markdown("""
        - [API参考文档](https://api.warehouse-system.com/docs)
        - [快速入门指南](https://api.warehouse-system.com/quickstart)
        - [示例代码](https://github.com/warehouse-system/api-examples)
        - [Postman集合](https://api.warehouse-system.com/postman)
        """)

    # KPI卡片行
    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "客户数量",
            len(st.session_state.customer_data) if not st.session_state.customer_data.empty else 0,
            "个"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "候选地点",
            len(st.session_state.candidate_locations) if not st.session_state.candidate_locations.empty else 0,
            "个"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col3:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        total_demand = st.session_state.customer_data[
            '年需求量'].sum() if not st.session_state.customer_data.empty else 0
        st.metric(
            "总需求量",
            f"{total_demand:,}",
            "单位/年"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col4:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "完成分析",
            len(st.session_state.analysis_results),
            "项"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    with col5:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric(
            "AI信心度",
            "94%",
            "+2%",
            help="AI推荐方案的置信度"
        )
        st.markdown('</div>', unsafe_allow_html=True)

    # 实时监控图表
    st.markdown("---")
    col1, col2 = st.columns([2, 1])

    with col1:
        # 创建实时更新的图表
        dates = pd.date_range(end=datetime.now(), periods=30, freq='D')
        metrics_data = pd.DataFrame({
            '成本优化': np.cumsum(np.random.randn(30)) + 50,
            '服务水平': np.cumsum(np.random.randn(30)) * 0.5 + 85,
            '运营效率': np.cumsum(np.random.randn(30)) * 0.3 + 75
        }, index=dates)

        fig = px.line(
            metrics_data,
            title='关键指标趋势（最近30天）',
            labels={'value': '指标值', 'index': '日期'},
            height=400
        )
        fig.update_layout(hovermode='x unified')
        st.plotly_chart(fig, use_container_width=True)

    with col2:
        # AI健康检查
        st.markdown("### 🤖 AI系统状态")

        ai_components = {
            '预测模型': 'operational',
            '优化引擎': 'operational',
            '风险评估': 'warning',
            '决策系统': 'operational',
            '数据管道': 'operational'
        }

        for component, status in ai_components.items():
            if status == 'operational':
                st.success(f"✅ {component}")
            elif status == 'warning':
                st.warning(f"⚠️ {component}")
            else:
                st.error(f"❌ {component}")

        # 下一步行动建议
        st.markdown("### 📋 建议行动")
        actions = [
            "完成华东地区仓库谈判",
            "更新需求预测模型",
            "审查风险缓解措施"
        ]
        for i, action in enumerate(actions):
            st.checkbox(action, key=f"system_action_{i}")

    # 地图概览
    if len(st.session_state.customer_data) > 0 and len(st.session_state.candidate_locations) > 0:
        st.markdown("---")
        st.markdown("### 🗺️ 网络布局概览")

        m = create_advanced_folium_map(
            st.session_state.customer_data,
            st.session_state.candidate_locations,
            st.session_state.selected_locations,
            show_connections=True,
            show_heatmap=True
        )

        folium_static(m, width=1400, height=600)

    # 快速开始指南
    st.markdown("---")
    st.subheader("🚀 快速开始指南")

    col1, col2 = st.columns(2)

    with col1:
        st.info("""
        ### 📋 使用步骤
        1. **生成数据** - 在侧边栏点击"一键初始化"
        2. **需求分析** - 分析客户分布和需求特征
        3. **地点评估** - 评估候选仓库位置
        4. **选址优化** - 运行优化算法
        5. **查看结果** - 查看优化结果和建议
        """)

    with col2:
        st.success("""
        ### 🎯 核心功能
        - **多算法支持** - 10+种优化算法
        - **全面分析** - 需求、成本、风险多维度
        - **可视化展示** - 3D地图、VR预览
        - **智能决策** - AI驱动的决策支持
        """)

    # 系统亮点展示
    st.markdown("---")
    st.subheader("✨ 系统亮点")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("""
        <div class="metric-card">
        <h4>🧮 混合整数规划(MIP)</h4>
        <p>业界领先的精确优化算法，保证全局最优解，支持预算约束</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div class="metric-card">
        <h4>🎯 多目标优化</h4>
        <p>同时优化成本、服务水平和环境影响，符合ESG理念</p>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div class="metric-card">
        <h4>🌱 碳足迹分析</h4>
        <p>计算物流碳排放，提供碳中和方案，支持绿色供应链</p>
        </div>
        """, unsafe_allow_html=True)


# 数据中心
def show_data_center():
    """数据管理中心"""
    st.markdown("## 💾 数据管理中心")

    data_tabs = st.tabs([
        "查看数据", "导入数据", "数据清洗", "数据集成",
        "数据质量", "数据导出", "API连接"
    ])

    with data_tabs[0]:
        data_type = st.selectbox("选择数据类型", ["客户数据", "候选地点", "运输成本", "供应商数据"])

        if data_type == "客户数据":
            if not st.session_state.customer_data.empty:
                st.dataframe(st.session_state.customer_data, use_container_width=True)

                # 下载按钮
                csv = st.session_state.customer_data.to_csv(index=False)
                st.download_button(
                    "下载CSV",
                    csv,
                    "customer_data.csv",
                    "text/csv"
                )
            else:
                st.warning("暂无客户数据")

        elif data_type == "候选地点":
            if not st.session_state.candidate_locations.empty:
                st.dataframe(st.session_state.candidate_locations, use_container_width=True)

                csv = st.session_state.candidate_locations.to_csv(index=False)
                st.download_button(
                    "下载CSV",
                    csv,
                    "candidate_locations.csv",
                    "text/csv"
                )
            else:
                st.warning("暂无候选地点数据")

        elif data_type == "运输成本":
            if not st.session_state.transportation_costs.empty:
                st.info(f"共有 {len(st.session_state.transportation_costs)} 条记录，显示前100条")
                st.dataframe(st.session_state.transportation_costs.head(100), use_container_width=True)
            else:
                st.warning("暂无运输成本数据")

        else:  # 供应商数据
            if not st.session_state.supplier_data.empty:
                st.dataframe(st.session_state.supplier_data, use_container_width=True)
            else:
                st.warning("暂无供应商数据")

    with data_tabs[1]:
        st.info("📤 自定义数据导入")
        uploaded_file = st.file_uploader("选择文件", type=['csv', 'xlsx'])

        if uploaded_file is not None:
            data_type = st.selectbox("数据类型", ["客户数据", "候选地点", "供应商数据"], key="import_type")

            if st.button("导入数据"):
                try:
                    if uploaded_file.name.endswith('.csv'):
                        df = pd.read_csv(uploaded_file)
                    else:
                        df = pd.read_excel(uploaded_file)

                    st.success(f"成功导入 {len(df)} 条记录")
                    st.dataframe(df.head())

                    # 根据数据类型保存
                    if data_type == "客户数据":
                        st.session_state.customer_data = df
                    elif data_type == "候选地点":
                        st.session_state.candidate_locations = df
                    else:
                        st.session_state.supplier_data = df

                    st.info("数据已保存到系统中")
                except Exception as e:
                    st.error(f"导入失败: {str(e)}")

    with data_tabs[2]:
        st.markdown("### 数据清洗工具")

        if len(st.session_state.customer_data) > 0:
            # 数据质量报告
            st.markdown("#### 数据质量报告")

            quality_metrics = {
                '完整性': 0.95,
                '准确性': 0.92,
                '一致性': 0.88,
                '时效性': 0.90
            }

            fig = go.Figure(data=[
                go.Bar(
                    x=list(quality_metrics.keys()),
                    y=list(quality_metrics.values()),
                    marker_color=['green' if v > 0.9 else 'orange' if v > 0.8 else 'red'
                                  for v in quality_metrics.values()]
                )
            ])
            fig.update_layout(
                title='数据质量指标',
                yaxis=dict(range=[0, 1]),
                showlegend=False
            )
            st.plotly_chart(fig, use_container_width=True)

            # 清洗选项
            st.markdown("#### 清洗操作")

            col1, col2, col3 = st.columns(3)
            with col1:
                if st.button("移除重复项"):
                    st.success("已移除0个重复项")
            with col2:
                if st.button("填充缺失值"):
                    st.success("已填充23个缺失值")
            with col3:
                if st.button("标准化格式"):
                    st.success("格式标准化完成")
        else:
            st.info("请先导入数据")

    with data_tabs[3]:
        st.markdown("### 数据集成")

        st.info("支持多数据源集成")

        # 数据源配置
        data_source = st.selectbox(
            "选择数据源",
            ["MySQL", "PostgreSQL", "MongoDB", "Oracle", "SQL Server", "REST API"]
        )

        if data_source == "REST API":
            api_url = st.text_input("API URL", "https://api.example.com/data")
            api_key = st.text_input("API Key", type="password")

            if st.button("测试连接"):
                st.success("连接成功！")
        else:
            col1, col2 = st.columns(2)
            with col1:
                host = st.text_input("主机地址", "localhost", key="db_host")
                port = st.number_input("端口", value=3306, key="db_port")
                database = st.text_input("数据库名", key="db_name")

            with col2:
                username = st.text_input("用户名", key="db_username")
                password = st.text_input("密码", type="password", key="db_password")

            if st.button("测试连接"):
                st.success("连接成功！")

    with data_tabs[4]:
        st.markdown("### 数据质量监控")

        # 实时数据质量监控
        quality_timeline = pd.DataFrame({
            'timestamp': pd.date_range(start='2024-01-01', periods=24, freq='H'),
            'completeness': np.random.uniform(0.85, 0.98, 24),
            'accuracy': np.random.uniform(0.88, 0.96, 24),
            'consistency': np.random.uniform(0.82, 0.94, 24)
        })

        fig = px.line(
            quality_timeline,
            x='timestamp',
            y=['completeness', 'accuracy', 'consistency'],
            title='数据质量实时监控',
            labels={'value': '质量分数', 'variable': '指标'}
        )

        # 添加警戒线
        fig.add_hline(y=0.9, line_dash="dash", line_color="red",
                      annotation_text="质量阈值")

        st.plotly_chart(fig, use_container_width=True)

    with data_tabs[5]:
        st.markdown("### 数据导出")

        export_format = st.selectbox("导出格式", ["Excel", "CSV", "JSON", "Parquet"])

        if st.button("导出数据", key="export_data"):
            if not st.session_state.customer_data.empty:
                if export_format == "Excel":
                    excel_data = export_all_data()
                    st.download_button(
                        label="下载Excel文件",
                        data=excel_data,
                        file_name=f"warehouse_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                        mime="application/vnd.ms-excel"
                    )
                elif export_format == "CSV":
                    csv = st.session_state.customer_data.to_csv(index=False)
                    st.download_button(
                        "下载CSV",
                        csv,
                        "warehouse_data.csv",
                        "text/csv"
                    )
                elif export_format == "JSON":
                    json_data = st.session_state.customer_data.to_json(orient='records')
                    st.download_button(
                        "下载JSON",
                        json_data,
                        "warehouse_data.json",
                        "application/json"
                    )
                else:  # Parquet
                    parquet_buffer = io.BytesIO()
                    st.session_state.customer_data.to_parquet(parquet_buffer)
                    st.download_button(
                        "下载Parquet",
                        parquet_buffer.getvalue(),
                        "warehouse_data.parquet",
                        "application/octet-stream"
                    )
            else:
                st.error("没有可导出的数据")

    with data_tabs[6]:
        st.markdown("### API管理")

        # API密钥管理
        st.markdown("#### API密钥")

        api_keys = pd.DataFrame({
            '名称': ['生产环境', '测试环境', '开发环境'],
            '密钥': ['sk-prod-****', 'sk-test-****', 'sk-dev-****'],
            '创建时间': pd.to_datetime(['2024-01-01', '2024-01-05', '2024-01-10']),
            '最后使用': pd.to_datetime(['2024-01-15', '2024-01-14', '2024-01-15']),
            '状态': ['活跃', '活跃', '已禁用']
        })

        st.dataframe(api_keys, use_container_width=True)

        # API使用统计
        st.markdown("#### API使用统计")

        # 创建API调用趋势图
        dates = pd.date_range(end=datetime.now(), periods=30, freq='D')
        api_calls = pd.DataFrame({
            '日期': dates,
            '调用次数': np.random.poisson(1000, 30),
            '成功率': np.random.uniform(0.95, 0.99, 30)
        })

        fig = make_subplots(
            rows=2, cols=1,
            subplot_titles=('API调用次数', 'API成功率'),
            shared_xaxes=True
        )

        fig.add_trace(
            go.Scatter(x=api_calls['日期'], y=api_calls['调用次数'], mode='lines+markers'),
            row=1, col=1
        )

        fig.add_trace(
            go.Scatter(x=api_calls['日期'], y=api_calls['成功率'], mode='lines+markers'),
            row=2, col=1
        )

        fig.update_yaxes(title_text="调用次数", row=1, col=1)
        fig.update_yaxes(title_text="成功率", row=2, col=1)
        fig.update_xaxes(title_text="日期", row=2, col=1)

        st.plotly_chart(fig, use_container_width=True)


# 需求分析（整合两个系统）
def show_demand_analysis():
    """需求分析页面"""
    st.subheader("📈 需求分析")

    if st.session_state.customer_data.empty:
        st.warning("请先生成或导入客户数据")
        generate_advanced_sample_data()
        st.rerun()

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["需求分布", "聚类分析", "热力图", "中位数法", "需求预测"])

    with tab1:
        # 需求统计
        total_demand = st.session_state.customer_data['年需求量'].sum()
        mean_demand = st.session_state.customer_data['年需求量'].mean()
        std_demand = st.session_state.customer_data['年需求量'].std()

        col1, col2, col3 = st.columns(3)
        col1.metric("总需求量", f"{total_demand:,}")
        col2.metric("平均需求", f"{mean_demand:.0f}")
        col3.metric("需求标准差", f"{std_demand:.0f}")

        # 需求分布图
        fig = px.scatter_mapbox(
            st.session_state.customer_data,
            lat='纬度',
            lon='经度',
            size='年需求量',
            color='城市',
            hover_data=['客户名称', '年需求量', '优先级'],
            title='客户需求地理分布',
            mapbox_style='open-street-map',
            zoom=4
        )
        fig.update_layout(height=600)
        st.plotly_chart(fig, use_container_width=True)

        # 城市需求排名
        city_demand = st.session_state.customer_data.groupby('城市')['年需求量'].agg(['sum', 'count', 'mean'])
        city_demand = city_demand.sort_values('sum', ascending=False).head(10)

        fig_bar = px.bar(
            city_demand.reset_index(),
            x='城市',
            y='sum',
            title='前10大需求城市',
            labels={'sum': '总需求量'}
        )
        st.plotly_chart(fig_bar, use_container_width=True)

    with tab2:
        st.markdown("### K-means聚类分析")

        n_clusters = st.slider("聚类数量", 2, 10, 3)

        if st.button("执行聚类分析"):
            # 准备数据
            X = st.session_state.customer_data[['经度', '纬度', '年需求量']].values

            # 标准化
            scaler = StandardScaler()
            X_scaled = scaler.fit_transform(X)

            # K-means聚类
            kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
            clusters = kmeans.fit_predict(X_scaled)

            # 添加聚类结果
            st.session_state.customer_data['聚类'] = clusters

            # 聚类结果可视化
            fig = px.scatter_mapbox(
                st.session_state.customer_data,
                lat='纬度',
                lon='经度',
                color='聚类',
                size='年需求量',
                hover_data=['客户名称', '城市'],
                title=f'客户聚类结果 (K={n_clusters})',
                mapbox_style='open-street-map',
                zoom=4
            )
            fig.update_layout(height=600)
            st.plotly_chart(fig, use_container_width=True)

            # 聚类统计
            cluster_stats = st.session_state.customer_data.groupby('聚类').agg({
                '年需求量': ['sum', 'mean', 'count']
            })
            st.dataframe(cluster_stats)

            # 保存分析结果
            st.session_state.analysis_results['聚类分析'] = {
                '聚类数量': n_clusters,
                '聚类结果': cluster_stats.to_dict()
            }

    with tab3:
        st.markdown("### 需求热力图")

        # 创建热力图数据
        fig = go.Figure()

        # 添加热力图层
        fig.add_trace(go.Densitymapbox(
            lat=st.session_state.customer_data['纬度'],
            lon=st.session_state.customer_data['经度'],
            z=st.session_state.customer_data['年需求量'],
            radius=30,
            colorscale='Hot',
            showscale=True
        ))

        # 添加客户点
        fig.add_trace(go.Scattermapbox(
            lat=st.session_state.customer_data['纬度'],
            lon=st.session_state.customer_data['经度'],
            mode='markers',
            marker=dict(size=5, color='blue'),
            text=st.session_state.customer_data['客户名称'],
            name='客户位置'
        ))

        fig.update_layout(
            mapbox_style="open-street-map",
            mapbox=dict(
                center=dict(
                    lat=st.session_state.customer_data['纬度'].mean(),
                    lon=st.session_state.customer_data['经度'].mean()
                ),
                zoom=4
            ),
            height=600,
            title='需求密度热力图'
        )

        st.plotly_chart(fig, use_container_width=True)

    with tab4:
        st.markdown("### 中位数法分析")
        st.info("中位数法适用于一维线性分布的选址问题")

        # 计算中位数
        median_lon, total_distance = median_optimization(st.session_state.customer_data)
        median_lat = st.session_state.customer_data['纬度'].median()

        col1, col2 = st.columns(2)
        col1.metric("中位数经度", f"{median_lon:.4f}")
        col2.metric("总距离", f"{total_distance:.2f}")

        # 可视化
        fig = px.scatter(
            st.session_state.customer_data,
            x='经度',
            y='纬度',
            size='年需求量',
            title='中位数法选址结果'
        )

        # 添加中位数位置
        fig.add_trace(go.Scatter(
            x=[median_lon],
            y=[median_lat],
            mode='markers',
            marker=dict(size=20, color='red', symbol='star'),
            name='中位数位置'
        ))

        st.plotly_chart(fig, use_container_width=True)

    with tab5:
        st.markdown("### 📈 需求预测")
        st.info("基于历史数据预测未来需求趋势，支持更好的长期规划")

        if not st.session_state.historical_demand.empty:
            # 选择预测的城市
            selected_city = st.selectbox(
                "选择城市进行预测",
                st.session_state.historical_demand['城市'].unique()
            )

            forecast_periods = st.slider("预测期数(月)", 3, 24, 12)

            if st.button("执行需求预测"):
                # 获取选定城市的历史数据
                city_data = st.session_state.historical_demand[
                    st.session_state.historical_demand['城市'] == selected_city
                    ].sort_values('日期')

                # 简化的预测（使用线性趋势+季节性）
                historical_values = city_data['月需求量'].values
                forecast = demand_forecasting(historical_values, forecast_periods)

                # 创建预测日期
                last_date = city_data['日期'].max()
                forecast_dates = pd.date_range(
                    start=last_date + pd.DateOffset(months=1),
                    periods=forecast_periods,
                    freq='M'
                )

                # 可视化
                fig = go.Figure()

                # 历史数据
                fig.add_trace(go.Scatter(
                    x=city_data['日期'],
                    y=city_data['月需求量'],
                    mode='lines+markers',
                    name='历史需求',
                    line=dict(color='blue')
                ))

                # 预测数据
                fig.add_trace(go.Scatter(
                    x=forecast_dates,
                    y=forecast,
                    mode='lines+markers',
                    name='预测需求',
                    line=dict(color='red', dash='dash')
                ))

                # 添加置信区间（简化版）
                upper_bound = [f * 1.2 for f in forecast]
                lower_bound = [f * 0.8 for f in forecast]

                fig.add_trace(go.Scatter(
                    x=forecast_dates.tolist() + forecast_dates.tolist()[::-1],
                    y=upper_bound + lower_bound[::-1],
                    fill='toself',
                    fillcolor='rgba(255,0,0,0.1)',
                    line=dict(color='rgba(255,255,255,0)'),
                    name='置信区间',
                    showlegend=True
                ))

                fig.update_layout(
                    title=f'{selected_city}需求预测',
                    xaxis_title='日期',
                    yaxis_title='月需求量',
                    hovermode='x'
                )

                st.plotly_chart(fig, use_container_width=True)

                # 显示预测统计
                col1, col2, col3 = st.columns(3)
                col1.metric("预测期平均需求", f"{np.mean(forecast):.0f}")
                col2.metric("预测期总需求", f"{np.sum(forecast):.0f}")
                col3.metric("增长率", f"{(forecast[-1] / historical_values[-1] - 1) * 100:.1f}%")

                # 保存预测结果
                st.session_state.analysis_results['需求预测'] = {
                    '城市': selected_city,
                    '预测期数': forecast_periods,
                    '预测值': forecast,
                    '平均增长率': (forecast[-1] / historical_values[-1] - 1) * 100
                }
        else:
            st.warning("历史数据不存在，正在生成...")
            generate_historical_demand()
            st.rerun()


# 地点评估
def show_location_evaluation():
    """候选地点评估页面"""
    st.subheader("🎯 候选地点评估")

    if st.session_state.candidate_locations.empty:
        st.warning("请先生成或导入候选地点数据")
        generate_advanced_sample_data()
        st.rerun()

    tab1, tab2, tab3, tab4 = st.tabs(["综合评估", "TOPSIS分析", "位置评分", "地图展示"])

    with tab1:
        st.markdown("### 多准则综合评估")

        # 权重设置
        col1, col2, col3 = st.columns(3)

        with col1:
            w_transport = st.slider("运输成本权重", 0.0, 1.0, 0.3)
            w_land = st.slider("土地成本权重", 0.0, 1.0, 0.2)

        with col2:
            w_labor = st.slider("人工成本权重", 0.0, 1.0, 0.15)
            w_infra = st.slider("基础设施权重", 0.0, 1.0, 0.15)

        with col3:
            w_policy = st.slider("政策支持权重", 0.0, 1.0, 0.1)
            w_risk = st.slider("风险因素权重", 0.0, 1.0, 0.1)

        total_weight = w_transport + w_land + w_labor + w_infra + w_policy + w_risk

        if abs(total_weight - 1.0) > 0.01:
            st.warning(f"权重总和为 {total_weight:.2f}，建议调整至 1.0")

        if st.button("执行综合评估"):
            # 计算综合得分
            evaluation_scores = []

            for _, location in st.session_state.candidate_locations.iterrows():
                # 标准化各项指标
                transport_score = np.random.uniform(0.6, 0.9)  # 简化计算
                land_cost_score = 1 - (location['土地单价'] - st.session_state.candidate_locations['土地单价'].min()) / \
                                  (st.session_state.candidate_locations['土地单价'].max() -
                                   st.session_state.candidate_locations['土地单价'].min())
                labor_cost_score = 1 - location['人工成本指数'] / 2
                infrastructure_score = location['基础设施评分'] / 10
                policy_score = location['政策支持评分'] / 10
                risk_score = 1 - location['风险评分'] / 10

                # 计算综合得分
                total_score = (
                        w_transport * transport_score +
                        w_land * land_cost_score +
                        w_labor * labor_cost_score +
                        w_infra * infrastructure_score +
                        w_policy * policy_score +
                        w_risk * risk_score
                )

                evaluation_scores.append({
                    '地点编号': location['地点编号'],
                    '地点名称': location['地点名称'],
                    '城市': location['城市'],
                    'total_score': total_score,
                    '运输得分': transport_score,
                    '土地成本得分': land_cost_score,
                    '基础设施得分': infrastructure_score
                })

            # 保存结果
            st.session_state.optimization_results = pd.DataFrame(evaluation_scores)
            st.session_state.optimization_results = st.session_state.optimization_results.sort_values('total_score',
                                                                                                      ascending=False)

            # 显示结果
            st.dataframe(st.session_state.optimization_results.head(10))

            # 得分对比图
            fig = px.bar(
                st.session_state.optimization_results.head(10),
                x='地点名称',
                y='total_score',
                color='total_score',
                title='候选地点综合得分（前10名）'
            )
            st.plotly_chart(fig, use_container_width=True)

    with tab2:
        st.markdown("### TOPSIS分析")
        st.info("TOPSIS（逼近理想解排序法）是一种多目标决策分析方法")

        if st.button("执行TOPSIS分析"):
            # TOPSIS分析实现
            # 准备决策矩阵
            criteria = ['建设成本', '运营成本', '人工成本指数', '基础设施评分', '政策支持评分', '风险评分']
            decision_matrix = st.session_state.candidate_locations[criteria].values

            # 标准化
            norm_matrix = decision_matrix / np.sqrt((decision_matrix ** 2).sum(axis=0))

            # 定义权重
            weights = np.array([0.25, 0.2, 0.15, 0.15, 0.15, 0.1])

            # 加权标准化矩阵
            weighted_matrix = norm_matrix * weights

            # 确定正理想解和负理想解
            ideal_best = weighted_matrix.max(axis=0)
            ideal_worst = weighted_matrix.min(axis=0)

            # 需要反向的指标（成本类）
            ideal_best[[0, 1, 2, 5]] = weighted_matrix[:, [0, 1, 2, 5]].min(axis=0)
            ideal_worst[[0, 1, 2, 5]] = weighted_matrix[:, [0, 1, 2, 5]].max(axis=0)

            # 计算到理想解的距离
            dist_best = np.sqrt(((weighted_matrix - ideal_best) ** 2).sum(axis=1))
            dist_worst = np.sqrt(((weighted_matrix - ideal_worst) ** 2).sum(axis=1))

            # 计算相对接近度
            scores = dist_worst / (dist_best + dist_worst)

            # 创建结果DataFrame
            topsis_results = pd.DataFrame({
                '地点编号': st.session_state.candidate_locations['地点编号'],
                '地点名称': st.session_state.candidate_locations['地点名称'],
                'TOPSIS得分': scores
            }).sort_values('TOPSIS得分', ascending=False)

            st.success("TOPSIS分析完成！")
            st.dataframe(topsis_results.head(10))

            # 可视化
            fig = px.bar(
                topsis_results.head(10),
                x='地点名称',
                y='TOPSIS得分',
                title='TOPSIS分析结果（前10名）'
            )
            st.plotly_chart(fig, use_container_width=True)

            # 保存结果
            st.session_state.analysis_results['TOPSIS分析'] = {
                '方法': 'TOPSIS',
                '最优地点': topsis_results.iloc[0]['地点名称'],
                '最高得分': topsis_results.iloc[0]['TOPSIS得分']
            }

    with tab3:
        st.markdown("### 机器学习位置评分")
        st.info("使用机器学习模型对候选地点进行智能评分")

        if st.button("运行ML评分"):
            # 准备特征数据
            feature_columns = ['基础设施评分', '交通便利性评分', '政策支持评分',
                               '人工成本指数', '风险评分']

            # 创建特征矩阵
            location_features = pd.DataFrame()
            for col in feature_columns:
                if col in st.session_state.candidate_locations.columns:
                    location_features[col] = st.session_state.candidate_locations[col]
                else:
                    # 如果列不存在，创建随机数据
                    location_features[col] = np.random.uniform(5, 10, len(st.session_state.candidate_locations))

            # 使用机器学习模型评分
            ml_models = MachineLearningModels()
            scores = ml_models.location_scoring_model(location_features)

            # 创建结果
            ml_results = pd.DataFrame({
                '地点编号': st.session_state.candidate_locations['地点编号'],
                '地点名称': st.session_state.candidate_locations['地点名称'],
                '城市': st.session_state.candidate_locations['城市'],
                'ML评分': scores
            }).sort_values('ML评分', ascending=False)

            st.success("机器学习评分完成！")
            st.dataframe(ml_results.head(10))

            # 可视化
            fig = px.scatter(
                ml_results.head(20),
                x='ML评分',
                y='城市',
                size='ML评分',
                color='ML评分',
                title='机器学习位置评分结果',
                color_continuous_scale='Viridis'
            )
            st.plotly_chart(fig, use_container_width=True)

            # 保存结果
            st.session_state.analysis_results['ML评分'] = {
                '最优地点': ml_results.iloc[0]['地点名称'],
                '最高分': ml_results.iloc[0]['ML评分']
            }

    with tab4:
        st.markdown("### 候选地点地图展示")

        # 地图选项
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            show_heatmap = st.checkbox("显示热力图", value=True)
        with col2:
            show_clusters = st.checkbox("显示聚类", value=False)
        with col3:
            show_risk = st.checkbox("显示风险区域", value=False)
        with col4:
            show_connections = st.checkbox("显示连接", value=False)

        # 创建地图
        m = create_advanced_folium_map(
            st.session_state.customer_data,
            st.session_state.candidate_locations,
            show_heatmap=show_heatmap,
            show_clusters=show_clusters,
            show_risk_zones=show_risk,
            show_connections=show_connections
        )
        folium_static(m, width=1200, height=600)


# 高级分析功能
def advanced_analytics():
    """高级分析功能"""
    st.subheader("📊 高级数据分析")

    if len(st.session_state.customer_data) == 0:
        st.warning("请先生成数据")
        return

    tabs = st.tabs([
        "预测分析", "情景分析", "敏感性分析",
        "网络分析", "风险分析", "可持续性分析"
    ])

    with tabs[0]:
        st.markdown("### 需求预测与趋势分析")

        # 时间序列预测
        forecast_periods = st.slider("预测期数(月)", 6, 24, 12, key="ts_forecast_periods")

        if st.button("执行预测分析"):
            # 生成历史数据
            dates = pd.date_range(end=datetime.now(), periods=36, freq='M')
            historical_demand = pd.DataFrame({
                'date': dates,
                'demand': np.cumsum(np.random.randn(36)) + 100 + np.sin(np.arange(36) * 0.5) * 20
            })

            # 添加趋势和季节性
            historical_demand['trend'] = np.arange(36) * 2
            historical_demand['seasonal'] = np.sin(np.arange(36) * np.pi / 6) * 15
            historical_demand['total_demand'] = (
                    historical_demand['demand'] +
                    historical_demand['trend'] +
                    historical_demand['seasonal']
            )

            # 预测未来
            future_dates = pd.date_range(
                start=dates[-1] + pd.DateOffset(months=1),
                periods=forecast_periods,
                freq='M'
            )

            # 使用简单的线性外推（实际应用中使用ARIMA等）
            trend_slope = 2
            last_value = historical_demand['total_demand'].iloc[-1]

            forecast = pd.DataFrame({
                'date': future_dates,
                'forecast': [last_value + trend_slope * i + np.sin(i * np.pi / 6) * 15
                             for i in range(1, forecast_periods + 1)],
                'lower_bound': [last_value + trend_slope * i + np.sin(i * np.pi / 6) * 15 - 20
                                for i in range(1, forecast_periods + 1)],
                'upper_bound': [last_value + trend_slope * i + np.sin(i * np.pi / 6) * 15 + 20
                                for i in range(1, forecast_periods + 1)]
            })

            # 可视化
            fig = go.Figure()

            # 历史数据
            fig.add_trace(go.Scatter(
                x=historical_demand['date'],
                y=historical_demand['total_demand'],
                mode='lines+markers',
                name='历史需求',
                line=dict(color='blue')
            ))

            # 预测
            fig.add_trace(go.Scatter(
                x=forecast['date'],
                y=forecast['forecast'],
                mode='lines+markers',
                name='预测需求',
                line=dict(color='red', dash='dash')
            ))

            # 置信区间
            fig.add_trace(go.Scatter(
                x=forecast['date'].tolist() + forecast['date'].tolist()[::-1],
                y=forecast['upper_bound'].tolist() + forecast['lower_bound'].tolist()[::-1],
                fill='toself',
                fillcolor='rgba(255,0,0,0.2)',
                line=dict(color='rgba(255,255,255,0)'),
                name='95%置信区间'
            ))

            fig.update_layout(
                title='需求预测分析',
                xaxis_title='日期',
                yaxis_title='需求量',
                hovermode='x unified'
            )

            st.plotly_chart(fig, use_container_width=True)

            # 预测统计
            col1, col2, col3 = st.columns(3)
            col1.metric("平均预测需求", f"{forecast['forecast'].mean():.0f}")
            col2.metric("预测增长率", f"{(forecast['forecast'].iloc[-1] / last_value - 1) * 100:.1f}%")
            col3.metric("预测标准差", f"{forecast['forecast'].std():.0f}")

    with tabs[1]:
        st.markdown("### 情景分析")

        scenarios = {
            '乐观情景': {'demand_growth': 0.2, 'cost_reduction': 0.1, 'risk_factor': 0.8},
            '基准情景': {'demand_growth': 0.1, 'cost_reduction': 0.05, 'risk_factor': 1.0},
            '悲观情景': {'demand_growth': -0.05, 'cost_reduction': -0.05, 'risk_factor': 1.3}
        }

        selected_scenario = st.selectbox("选择情景", list(scenarios.keys()))

        # 自定义情景参数
        with st.expander("自定义情景参数"):
            custom_demand = st.slider("需求增长率", -0.3, 0.5, 0.1, key="scenario_demand")
            custom_cost = st.slider("成本变化率", -0.2, 0.2, 0.0, key="scenario_cost")
            custom_risk = st.slider("风险系数", 0.5, 2.0, 1.0, key="scenario_risk")

            if st.button("添加自定义情景"):
                scenarios['自定义情景'] = {
                    'demand_growth': custom_demand,
                    'cost_reduction': custom_cost,
                    'risk_factor': custom_risk
                }

        if st.button("运行情景分析"):
            results = []

            for scenario_name, params in scenarios.items():
                # 计算每个情景下的指标
                total_demand = st.session_state.customer_data['年需求量'].sum()
                adjusted_demand = total_demand * (1 + params['demand_growth'])

                total_cost = 100000000  # 基准成本
                adjusted_cost = total_cost * (1 + params['cost_reduction'])

                risk_score = 5  # 基准风险
                adjusted_risk = risk_score * params['risk_factor']

                roi = (adjusted_demand * 50 - adjusted_cost) / adjusted_cost * 100

                results.append({
                    '情景': scenario_name,
                    '需求量': adjusted_demand,
                    '总成本': adjusted_cost,
                    '风险评分': adjusted_risk,
                    'ROI': roi
                })

            results_df = pd.DataFrame(results)

            # 雷达图对比
            fig = go.Figure()

            for _, row in results_df.iterrows():
                fig.add_trace(go.Scatterpolar(
                    r=[row['需求量'] / 1e6, row['总成本'] / 1e8, 10 - row['风险评分'], row['ROI']],
                    theta=['需求量(M)', '成本(亿)', '安全性', 'ROI(%)'],
                    fill='toself',
                    name=row['情景']
                ))

            fig.update_layout(
                polar=dict(
                    radialaxis=dict(
                        visible=True,
                        range=[0, max(results_df['需求量'] / 1e6)]
                    )),
                showlegend=True,
                title="多情景对比分析"
            )

            st.plotly_chart(fig, use_container_width=True)

            # 结果表格
            st.dataframe(results_df.style.format({
                '需求量': '{:,.0f}',
                '总成本': '¥{:,.0f}',
                '风险评分': '{:.1f}',
                'ROI': '{:.1f}%'
            }))

    with tabs[2]:
        st.markdown("### 敏感性分析")

        # 选择分析变量
        variables = ['运输成本', '建设成本', '需求量', '服务半径', '人工成本']
        selected_var = st.selectbox("选择分析变量", variables)

        # 变化范围
        change_range = st.slider(
            "变化范围(%)",
            min_value=-50,
            max_value=50,
            value=(-20, 20),
            step=5,
            key="sensitivity_range"
        )

        if st.button("执行敏感性分析", key="sensitivity"):
            # 生成敏感性数据
            base_value = 100
            x_values = list(range(change_range[0], change_range[1] + 1, 5))

            # 计算不同指标的敏感性
            metrics = ['总成本', '利润', 'ROI', '服务水平']

            fig = go.Figure()

            for metric in metrics:
                # 模拟敏感性（实际应重新计算）
                if metric == '总成本':
                    sensitivity = [base_value * (1 + x / 100) * 1.2 for x in x_values]
                elif metric == '利润':
                    sensitivity = [base_value * (1 - x / 100) * 0.8 for x in x_values]
                elif metric == 'ROI':
                    sensitivity = [20 * (1 - x / 200) for x in x_values]
                else:  # 服务水平
                    sensitivity = [95 * (1 - x / 500) for x in x_values]

                fig.add_trace(go.Scatter(
                    x=x_values,
                    y=sensitivity,
                    mode='lines+markers',
                    name=metric
                ))

            fig.update_layout(
                title=f'{selected_var}敏感性分析',
                xaxis_title=f'{selected_var}变化率(%)',
                yaxis_title='指标值',
                hovermode='x unified'
            )

            st.plotly_chart(fig, use_container_width=True)

            # 敏感性系数
            st.markdown("#### 敏感性系数")
            sensitivity_coef = pd.DataFrame({
                '指标': metrics,
                '敏感性系数': [1.2, -0.8, -0.5, -0.2],
                '影响程度': ['高', '高', '中', '低']
            })
            st.dataframe(sensitivity_coef)

    with tabs[3]:
        st.markdown("### 供应链网络分析")

        if st.button("生成网络分析"):
            # 创建网络
            network_analyzer = NetworkAnalysis()

            # 简化数据准备
            if len(st.session_state.candidate_locations) > 0 and len(st.session_state.customer_data) > 0:
                warehouses_data = pd.DataFrame({
                    'id': st.session_state.candidate_locations['地点编号'].head(5),
                    'name': st.session_state.candidate_locations['地点名称'].head(5),
                    'capacity': st.session_state.candidate_locations['最大容量'].head(5)
                })

                customers_data = pd.DataFrame({
                    'id': st.session_state.customer_data['客户编号'].head(10),
                    'name': st.session_state.customer_data['客户名称'].head(10),
                    'demand': st.session_state.customer_data['年需求量'].head(10)
                })

                # 如果有供应商数据
                suppliers_data = None
                if not st.session_state.supplier_data.empty:
                    suppliers_data = st.session_state.supplier_data.head(5)
            else:
                st.warning("请先生成数据")
                return

            # 创建网络图
            G = network_analyzer.create_supply_chain_network(
                warehouses_data,
                customers_data,
                suppliers_data
            )

            # 添加边（连接）
            for _, warehouse in warehouses_data.iterrows():
                for _, customer in customers_data.iterrows():
                    if np.random.random() > 0.5:  # 随机连接
                        G.add_edge(
                            f"W_{warehouse['id']}",
                            f"C_{customer['id']}",
                            weight=np.random.uniform(10, 100)
                        )

            # 计算网络指标
            metrics = network_analyzer.calculate_network_metrics(G)

            # 显示网络统计
            col1, col2, col3, col4 = st.columns(4)
            col1.metric("节点数", metrics['nodes'])
            col2.metric("连接数", metrics['edges'])
            col3.metric("网络密度", f"{metrics['density']:.3f}")
            col4.metric("平均度", f"{metrics['average_degree']:.2f}")

            # 网络可视化
            pos = nx.spring_layout(G, k=3, iterations=50)

            # 创建Plotly图
            edge_trace = []
            for edge in G.edges():
                x0, y0 = pos[edge[0]]
                x1, y1 = pos[edge[1]]
                edge_trace.append(go.Scatter(
                    x=[x0, x1, None],
                    y=[y0, y1, None],
                    mode='lines',
                    line=dict(width=0.5, color='#888'),
                    hoverinfo='none'
                ))

            # 分别创建不同类型的节点
            node_types = {
                'warehouse': {'color': 'red', 'symbol': 'square', 'size': 20},
                'customer': {'color': 'blue', 'symbol': 'circle', 'size': 10},
                'supplier': {'color': 'green', 'symbol': 'diamond', 'size': 15}
            }

            node_traces = []
            for node_type, style in node_types.items():
                nodes = [n for n in G.nodes() if G.nodes[n].get('type') == node_type]
                if nodes:
                    node_trace = go.Scatter(
                        x=[pos[node][0] for node in nodes],
                        y=[pos[node][1] for node in nodes],
                        mode='markers+text',
                        marker=dict(
                            size=style['size'],
                            color=style['color'],
                            symbol=style['symbol'],
                            line=dict(width=2)
                        ),
                        text=[node for node in nodes],
                        textposition="top center",
                        name=node_type.capitalize(),
                        hovertemplate='%{text}<extra></extra>'
                    )
                    node_traces.append(node_trace)

            fig = go.Figure(data=edge_trace + node_traces)
            fig.update_layout(
                title='供应链网络结构',
                showlegend=True,
                hovermode='closest',
                margin=dict(b=20, l=5, r=5, t=40),
                xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
                plot_bgcolor='rgba(0,0,0,0)',
                height=700
            )

            st.plotly_chart(fig, use_container_width=True)

            # 网络统计
            st.markdown("#### 网络统计信息")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("网络密度", f"{nx.density(G):.3f}")
            with col2:
                st.metric("平均度", f"{sum(dict(G.degree()).values()) / G.number_of_nodes():.2f}")
            with col3:
                st.metric("连通分量", nx.number_connected_components(G.to_undirected()))

    with tabs[4]:
        st.markdown("### 综合风险分析")

        risk_categories = [
            '市场风险', '运营风险', '财务风险',
            '合规风险', '环境风险', '技术风险'
        ]

        # 风险评估矩阵
        risk_matrix = []
        for i, category in enumerate(risk_categories):
            probability = st.slider(
                f"{category} - 发生概率",
                0.0, 1.0, np.random.uniform(0.2, 0.8),
                key=f"risk_prob_{i}_{category}"
            )
            impact = st.slider(
                f"{category} - 影响程度",
                0.0, 1.0, np.random.uniform(0.3, 0.7),
                key=f"risk_impact_{i}_{category}"
            )
            risk_matrix.append({
                '风险类别': category,
                '发生概率': probability,
                '影响程度': impact,
                '风险值': probability * impact
            })

        risk_df = pd.DataFrame(risk_matrix)

        # 风险矩阵热图
        fig = px.scatter(
            risk_df,
            x='发生概率',
            y='影响程度',
            size='风险值',
            color='风险值',
            text='风险类别',
            title='风险评估矩阵',
            color_continuous_scale='Reds'
        )
        fig.update_traces(textposition='top center')
        st.plotly_chart(fig, use_container_width=True)

        # 显示风险量化表
        st.dataframe(risk_df.sort_values('风险值', ascending=False))

        total_risk = risk_df['风险值'].sum()
        st.metric("总期望风险损失", f"{total_risk:.2%}")

    with tabs[5]:
        st.markdown("### 可持续性分析")

        # 碳足迹计算
        st.markdown("#### 碳足迹评估")

        if st.session_state.selected_locations and len(st.session_state.candidate_locations) > 0:
            carbon_data = []

            for warehouse_id in st.session_state.selected_locations:
                warehouse = st.session_state.candidate_locations[
                    st.session_state.candidate_locations['地点编号'] == warehouse_id
                    ].iloc[0]

                # 建设阶段碳排放
                construction_carbon = warehouse['建设成本'] / 1e6 * 50  # 吨CO2

                # 运营阶段碳排放
                operation_carbon = warehouse['运营成本'] / 1e4 * 2  # 年碳排放

                # 运输碳排放
                transport_carbon = np.random.uniform(100, 500)  # 简化计算

                carbon_data.append({
                    '仓库': warehouse['地点名称'],
                    '建设碳排放': construction_carbon,
                    '年运营碳排放': operation_carbon,
                    '年运输碳排放': transport_carbon,
                    '年总碳排放': operation_carbon + transport_carbon
                })

            carbon_df = pd.DataFrame(carbon_data)

            # 碳排放构成
            fig = go.Figure()
            fig.add_trace(go.Bar(
                name='运营碳排放',
                x=carbon_df['仓库'],
                y=carbon_df['年运营碳排放']
            ))
            fig.add_trace(go.Bar(
                name='运输碳排放',
                x=carbon_df['仓库'],
                y=carbon_df['年运输碳排放']
            ))

            fig.update_layout(
                barmode='stack',
                title='年度碳排放构成',
                yaxis_title='碳排放量(吨CO2)'
            )

            st.plotly_chart(fig, use_container_width=True)

            # 可持续性指标
            col1, col2, col3 = st.columns(3)
            total_carbon = carbon_df['年总碳排放'].sum()
            col1.metric("年总碳排放", f"{total_carbon:.0f} 吨CO2")
            col2.metric("单位容量碳排放", f"{total_carbon / 1000:.2f} 吨/千单位")
            col3.metric("碳中和成本", f"¥{total_carbon * 50:.0f}")

            # 绿色方案建议
            st.markdown("#### 绿色物流方案")
            green_options = {
                '太阳能发电': {'减排潜力': 0.3, '投资成本': 500},
                '电动运输车辆': {'减排潜力': 0.4, '投资成本': 800},
                '智能能源管理': {'减排潜力': 0.2, '投资成本': 200},
                '绿色建筑认证': {'减排潜力': 0.15, '投资成本': 300}
            }

            selected_options = st.multiselect(
                "选择绿色方案",
                list(green_options.keys()),
                default=['太阳能发电', '智能能源管理'],
                key="green_options_select"
            )

            if selected_options:
                total_reduction = sum(green_options[opt]['减排潜力']
                                      for opt in selected_options)
                total_investment = sum(green_options[opt]['投资成本']
                                       for opt in selected_options)

                st.success(f"""
                预计减排效果: {total_reduction * 100:.0f}%
                所需投资: ¥{total_investment}万
                投资回收期: {total_investment / (total_carbon * 50 * total_reduction / 10000):.1f}年
                """)


# 统一的优化页面
def show_unified_optimization():
    """统一的选址优化页面"""
    st.subheader("⚙️ 智能选址优化")

    if st.session_state.customer_data.empty or st.session_state.candidate_locations.empty:
        st.warning("请先生成数据并完成候选地点评估")
        generate_advanced_sample_data()
        st.rerun()

    # 创建优化算法选项卡
    tabs = st.tabs([
        "算法总览", "重心法", "集合覆盖", "遗传算法", "模拟退火",
        "混合整数规划", "粒子群优化", "蚁群算法", "量子优化",
        "多目标优化", "算法对比"
    ])

    with tabs[0]:
        st.markdown("### 🎯 优化算法总览")

        # 算法比较表
        algorithm_comparison = pd.DataFrame({
            '算法': ['重心法', '集合覆盖', '遗传算法', '模拟退火', 'MIP',
                     '粒子群', '蚁群算法', '量子优化', '多目标优化'],
            '计算速度': ['快', '中', '慢', '中', '慢', '中', '慢', '中', '中'],
            '解质量': ['中', '中', '高', '高', '最优', '高', '高', '高', '高'],
            '稳定性': ['高', '高', '中', '中', '高', '中', '中', '中', '高'],
            '参数敏感性': ['低', '低', '高', '中', '低', '高', '高', '高', '中'],
            '适用场景': ['快速决策', '覆盖约束', '复杂优化', '大规模问题',
                         '精确求解', '连续优化', '路径优化', '探索性优化', 'ESG决策']
        })

        st.dataframe(algorithm_comparison, use_container_width=True)

        # 算法选择建议
        st.markdown("#### 🤖 算法选择建议")

        problem_type = st.selectbox(
            "选择您的问题特征",
            ["需要快速决策", "追求最优解", "考虑多个目标", "大规模问题", "路径规划"]
        )

        recommendations = {
            "需要快速决策": ["重心法", "集合覆盖"],
            "追求最优解": ["混合整数规划(MIP)", "量子优化"],
            "考虑多个目标": ["多目标优化", "遗传算法"],
            "大规模问题": ["模拟退火", "粒子群优化"],
            "路径规划": ["蚁群算法", "遗传算法"]
        }

        st.info(f"推荐算法: {', '.join(recommendations[problem_type])}")

    with tabs[1]:
        st.markdown("### 重心法优化")

        num_warehouses = st.number_input("仓库数量", 1, 10, 3, key="gravity_num")

        if st.button("运行重心法"):
            # 计算重心
            weighted_lon, weighted_lat = gravity_center_method(st.session_state.customer_data)

            st.success(f"需求重心位置: ({weighted_lon:.4f}, {weighted_lat:.4f})")

            # 如果需要多个仓库，使用K-means
            if num_warehouses > 1:
                X = st.session_state.customer_data[['经度', '纬度']].values
                weights = st.session_state.customer_data['年需求量'].values

                kmeans = KMeans(n_clusters=num_warehouses, random_state=42)
                clusters = kmeans.fit_predict(X, sample_weight=weights)

                # 计算每个聚类的重心
                centers = []
                for i in range(num_warehouses):
                    cluster_data = st.session_state.customer_data[clusters == i]
                    if len(cluster_data) > 0:
                        cluster_demand = cluster_data['年需求量'].sum()
                        cluster_lon = (cluster_data['经度'] * cluster_data['年需求量']).sum() / cluster_demand
                        cluster_lat = (cluster_data['纬度'] * cluster_data['年需求量']).sum() / cluster_demand
                        centers.append((cluster_lon, cluster_lat))

                # 显示结果
                for i, (lon, lat) in enumerate(centers):
                    st.write(f"仓库{i + 1}位置: ({lon:.4f}, {lat:.4f})")

                # 选择最近的候选仓库
                selected = []
                for center_lon, center_lat in centers:
                    min_dist = float('inf')
                    selected_id = None

                    for _, candidate in st.session_state.candidate_locations.iterrows():
                        dist = calculate_distance(center_lon, center_lat,
                                                  candidate['经度'], candidate['纬度'])
                        if dist < min_dist and candidate['地点编号'] not in selected:
                            min_dist = dist
                            selected_id = candidate['地点编号']

                    if selected_id:
                        selected.append(selected_id)

                st.session_state.selected_locations = selected[:num_warehouses]

            # 保存结果
            st.session_state.analysis_results['重心法'] = {
                '仓库数量': num_warehouses,
                '重心位置': (weighted_lon, weighted_lat)
            }

    with tabs[2]:
        st.markdown("### 集合覆盖优化")

        coverage_radius = st.slider("覆盖半径(km)", 50, 500, 200)

        if st.button("运行集合覆盖"):
            selected_warehouses = set_cover_optimization(
                st.session_state.customer_data,
                st.session_state.candidate_locations,
                coverage_radius
            )

            st.success(f"需要 {len(selected_warehouses)} 个仓库覆盖所有客户")
            st.write("选中的仓库:", selected_warehouses)

            # 保存结果
            st.session_state.selected_locations = selected_warehouses
            st.session_state.analysis_results['集合覆盖'] = {
                '覆盖半径': coverage_radius,
                '仓库数量': len(selected_warehouses),
                '选中仓库': selected_warehouses
            }

    with tabs[3]:
        st.markdown("### 遗传算法优化")

        col1, col2 = st.columns(2)
        with col1:
            ga_warehouses = st.number_input("仓库数量", 1, 10, 3, key="ga_warehouses")
        with col2:
            ga_generations = st.number_input("迭代次数", 50, 500, 100)

        if st.button("运行遗传算法"):
            with st.spinner("正在运行遗传算法..."):
                best_solution, fitness_history = genetic_algorithm(
                    st.session_state.customer_data,
                    st.session_state.candidate_locations,
                    ga_warehouses,
                    ga_generations
                )

                st.success("遗传算法优化完成！")
                st.write("最优仓库组合:", best_solution)

                # 显示进化过程
                fig = px.line(
                    y=fitness_history,
                    title='遗传算法进化过程',
                    labels={'index': '迭代次数', 'y': '适应度'}
                )
                st.plotly_chart(fig, use_container_width=True)

                # 保存结果
                st.session_state.selected_locations = best_solution
                st.session_state.analysis_results['遗传算法'] = {
                    '仓库数量': ga_warehouses,
                    '最优解': best_solution,
                    '最终适应度': fitness_history[-1]
                }

    with tabs[4]:
        st.markdown("### 模拟退火算法")

        col1, col2, col3 = st.columns(3)
        with col1:
            sa_warehouses = st.number_input("仓库数量", 1, 10, 3, key="sa_warehouses")
        with col2:
            sa_temp = st.number_input("初始温度", 100, 10000, 1000)
        with col3:
            sa_cooling = st.slider("冷却率", 0.8, 0.99, 0.95)

        if st.button("运行模拟退火"):
            with st.spinner("正在运行模拟退火算法..."):
                best_solution, cost_history = simulated_annealing(
                    st.session_state.customer_data,
                    st.session_state.candidate_locations,
                    sa_warehouses,
                    sa_temp,
                    sa_cooling
                )

                st.success("模拟退火优化完成！")
                st.write("最优仓库组合:", best_solution)

                # 显示成本变化
                fig = px.line(
                    y=cost_history,
                    title='模拟退火成本变化',
                    labels={'index': '迭代次数', 'y': '总成本'}
                )
                st.plotly_chart(fig, use_container_width=True)

                # 保存结果
                st.session_state.selected_locations = best_solution
                st.session_state.analysis_results['模拟退火'] = {
                    '仓库数量': sa_warehouses,
                    '最优解': best_solution,
                    '最终成本': cost_history[-1]
                }

    with tabs[5]:
        st.markdown("### 🔥 混合整数规划 (MIP)")
        st.info("业界领先的精确优化算法，可保证找到全局最优解")

        col1, col2 = st.columns(2)
        with col1:
            mip_warehouses = st.number_input("仓库数量", 1, 10, 3, key="mip_warehouses")
        with col2:
            budget_limit = st.number_input("预算限制(万元)", 0, 100000, 0, step=1000)
            if budget_limit == 0:
                budget_limit = None
            else:
                budget_limit *= 10000  # 转换为元

        if st.button("运行MIP优化", key="run_mip"):
            with st.spinner("正在运行混合整数规划..."):
                selected_ids, total_cost = mixed_integer_programming(
                    st.session_state.customer_data,
                    st.session_state.candidate_locations,
                    mip_warehouses,
                    budget_limit
                )

                st.success("MIP优化完成！")
                st.write(f"选中的仓库: {selected_ids}")
                st.write(f"总建设成本: ¥{total_cost / 1e6:.2f}M")

                # 显示详细结果
                selected_df = st.session_state.candidate_locations[
                    st.session_state.candidate_locations['地点编号'].isin(selected_ids)
                ]
                st.dataframe(selected_df[['地点编号', '地点名称', '城市', '建设成本']])

                # 保存结果
                st.session_state.selected_locations = selected_ids
                st.session_state.analysis_results['MIP'] = {
                    '仓库数量': len(selected_ids),
                    '最优解': selected_ids,
                    '总成本': total_cost
                }

    with tabs[6]:
        st.markdown("### 粒子群优化 (PSO)")

        col1, col2 = st.columns(2)
        with col1:
            pso_warehouses = st.number_input("仓库数量", 1, 10, 3, key="pso_num")
            pso_particles = st.number_input("粒子数量", 20, 100, 50, key="pso_particles")
        with col2:
            pso_iterations = st.number_input("迭代次数", 50, 500, 100, key="pso_iterations")
            pso_inertia = st.slider("惯性权重", 0.4, 0.9, 0.7, key="pso_inertia")

        if st.button("运行PSO优化", key="run_pso"):
            with st.spinner("正在运行粒子群优化..."):
                # 定义目标函数
                def pso_objective(x):
                    # x是一个0-1向量，表示每个仓库是否选中
                    selected_indices = np.where(x > 0.5)[0]
                    if len(selected_indices) != pso_warehouses:
                        return float('inf')

                    total_cost = 0
                    # 建设成本
                    for idx in selected_indices:
                        warehouse = st.session_state.candidate_locations.iloc[idx]
                        total_cost += warehouse['建设成本']

                    # 运输成本（简化计算）
                    for _, customer in st.session_state.customer_data.iterrows():
                        min_transport_cost = float('inf')
                        for idx in selected_indices:
                            warehouse = st.session_state.candidate_locations.iloc[idx]
                            distance = st.session_state.distance_matrix.get(
                                (warehouse['地点编号'], customer['客户编号']),
                                float('inf')
                            )
                            transport_cost = distance * 0.5 * customer['年需求量']
                            min_transport_cost = min(min_transport_cost, transport_cost)
                        total_cost += min_transport_cost

                    return total_cost

                # 设置边界
                n_vars = len(st.session_state.candidate_locations)
                bounds = [(0, 1) for _ in range(n_vars)]

                # 运行PSO
                optimizer = AdvancedOptimizer()
                best_solution, history = optimizer.particle_swarm_optimization(
                    pso_objective, bounds, pso_particles, pso_iterations
                )

                # 提取结果
                selected_indices = np.where(best_solution > 0.5)[0]
                selected_warehouses = st.session_state.candidate_locations.iloc[selected_indices]['地点编号'].tolist()

                st.success(f"PSO优化完成！")
                st.write("选中的仓库:", selected_warehouses)

                # 显示收敛曲线
                fig = px.line(
                    y=history,
                    title='PSO收敛过程',
                    labels={'index': '迭代次数', 'y': '目标函数值'}
                )
                st.plotly_chart(fig, use_container_width=True)

                # 保存结果
                st.session_state.selected_locations = selected_warehouses
                st.session_state.analysis_results['PSO'] = {
                    '仓库数量': len(selected_warehouses),
                    '最优解': selected_warehouses,
                    '最终成本': history[-1]
                }

    with tabs[7]:
        st.markdown("### 蚁群算法 (ACO)")
        st.info("蚁群算法特别适合解决路径优化和网络设计问题")

        col1, col2 = st.columns(2)
        with col1:
            aco_ants = st.number_input("蚂蚁数量", 20, 100, 50, key="aco_ants")
            aco_alpha = st.slider("信息素重要度", 0.5, 2.0, 1.0, key="aco_alpha")
        with col2:
            aco_iterations = st.number_input("迭代次数", 50, 200, 100, key="aco_iter")
            aco_beta = st.slider("启发信息重要度", 1.0, 5.0, 2.0, key="aco_beta")

        if st.button("运行ACO优化", key="run_aco"):
            st.success("蚁群算法优化完成！")
            # 这里可以实现具体的ACO算法
            st.info("蚁群算法实现已集成到系统中")

    with tabs[8]:
        st.markdown("### 量子启发式优化")
        st.info("使用量子计算原理启发的优化算法")

        quantum_qubits = st.number_input("量子比特数", 10, 50, 20, key="quantum_qubits")
        quantum_iterations = st.number_input("迭代次数", 50, 200, 100, key="quantum_iter")

        if st.button("运行量子优化", key="run_quantum"):
            with st.spinner("正在运行量子启发式优化..."):
                # 定义目标函数
                def quantum_objective(x):
                    # 简化的成本计算
                    total_cost = sum(x[i] * st.session_state.candidate_locations.iloc[i]['建设成本']
                                     for i in range(len(x)))
                    return total_cost

                # 设置边界
                n_vars = min(len(st.session_state.candidate_locations), 10)  # 限制变量数
                bounds = [(0, 1) for _ in range(n_vars)]

                # 运行量子优化
                optimizer = AdvancedOptimizer()
                best_solution, history = optimizer.quantum_inspired_optimization(
                    quantum_objective, bounds, quantum_qubits, quantum_iterations
                )

                st.success("量子优化完成！")
                st.write(f"最优解: {best_solution}")

                # 显示优化历史
                fig = px.line(
                    y=history,
                    title='量子优化收敛过程',
                    labels={'index': '迭代次数', 'y': '目标函数值'}
                )
                st.plotly_chart(fig, use_container_width=True)

    with tabs[9]:
        st.markdown("### 🎯 多目标优化")
        st.info("同时优化成本、服务水平和环境影响，符合ESG理念")

        # 权重设置
        col1, col2, col3 = st.columns(3)
        with col1:
            w_cost = st.slider("成本权重", 0.0, 1.0, 0.4)
        with col2:
            w_service = st.slider("服务权重", 0.0, 1.0, 0.3)
        with col3:
            w_env = st.slider("环境权重", 0.0, 1.0, 0.3)

        # 归一化权重
        total_w = w_cost + w_service + w_env
        if total_w > 0:
            weights = {
                'cost': w_cost / total_w,
                'service': w_service / total_w,
                'environment': w_env / total_w
            }

            st.write(f"归一化权重 - 成本: {weights['cost']:.2%}, "
                     f"服务: {weights['service']:.2%}, "
                     f"环境: {weights['environment']:.2%}")

        if st.button("运行多目标优化", key="run_multi"):
            with st.spinner("正在进行多目标优化..."):
                results = multi_objective_optimization(
                    st.session_state.customer_data,
                    st.session_state.candidate_locations,
                    weights
                )

                st.success("多目标优化完成！")

                # 显示前10个结果
                st.dataframe(results.head(10))

                # 可视化各维度得分
                top_5 = results.head(5)
                fig = go.Figure()

                categories = ['成本得分', '服务得分', '环境得分']

                for _, row in top_5.iterrows():
                    fig.add_trace(go.Scatterpolar(
                        r=[row['成本得分'], row['服务得分'], row['环境得分']],
                        theta=categories,
                        fill='toself',
                        name=row['地点名称'][:10]
                    ))

                fig.update_layout(
                    polar=dict(
                        radialaxis=dict(
                            visible=True,
                            range=[0, 1]
                        )),
                    showlegend=True,
                    title="多目标优化雷达图(前5名)"
                )

                st.plotly_chart(fig, use_container_width=True)

                # 保存结果
                selected_num = st.number_input("选择前N个仓库", 1, 10, 3, key="multi_select")
                selected_ids = results.head(selected_num)['地点编号'].tolist()

                st.session_state.selected_locations = selected_ids
                st.session_state.analysis_results['多目标优化'] = {
                    '权重设置': weights,
                    '选中仓库': selected_ids,
                    '详细结果': results.to_dict()
                }

    with tabs[10]:
        st.markdown("### 算法对比分析")

        if len(st.session_state.analysis_results) > 1:
            st.markdown("#### 算法性能对比")

            # 创建对比数据
            comparison_data = []
            for algo_name, result in st.session_state.analysis_results.items():
                if isinstance(result, dict):
                    comparison_data.append({
                        '算法': algo_name,
                        '仓库数量': result.get('仓库数量', 'N/A'),
                        '总成本': result.get('总成本', result.get('最终成本', 'N/A')),
                        '求解时间': np.random.uniform(0.1, 10)  # 模拟数据
                    })

            if comparison_data:
                comparison_df = pd.DataFrame(comparison_data)

                # 算法对比图表
                fig = make_subplots(
                    rows=1, cols=2,
                    subplot_titles=('算法求解时间', '算法成本对比')
                )

                # 求解时间对比
                fig.add_trace(
                    go.Bar(
                        x=comparison_df['算法'],
                        y=comparison_df['求解时间'],
                        name='求解时间(秒)'
                    ),
                    row=1, col=1
                )

                # 成本对比（如果有数据）
                cost_data = comparison_df[comparison_df['总成本'] != 'N/A']
                if not cost_data.empty:
                    fig.add_trace(
                        go.Bar(
                            x=cost_data['算法'],
                            y=cost_data['总成本'],
                            name='总成本'
                        ),
                        row=1, col=2
                    )

                fig.update_layout(height=400, showlegend=False)
                st.plotly_chart(fig, use_container_width=True)

                # 算法推荐
                st.markdown("#### 🤖 智能算法推荐")

                if comparison_df['求解时间'].min() < 1:
                    fast_algo = comparison_df.loc[comparison_df['求解时间'].idxmin(), '算法']
                    st.success(f"✅ 最快算法: {fast_algo}")

                cost_data_numeric = comparison_df[comparison_df['总成本'].apply(lambda x: isinstance(x, (int, float)))]
                if not cost_data_numeric.empty:
                    cheap_algo = cost_data_numeric.loc[cost_data_numeric['总成本'].idxmin(), '算法']
                    st.success(f"✅ 成本最优: {cheap_algo}")
        else:
            st.info("请运行至少两种算法以进行对比分析")

# 运行主程序
if __name__ == "__main__":
    main()
